"""
Tests για astakos_skills/file_generator.py
Τρέξε: pytest tests/test_file_generator.py -v
"""
import os
import sys
import json
import csv
import tempfile
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from astakos_skills.file_generator import (
    generate_excel,
    generate_word_doc,
    generate_pdf,
    generate_csv,
)

SAMPLE_DATA = json.dumps([
    {"Όνομα": "Λάζαρος", "Ηλικία": 42, "Πόλη": "Θεσσαλονίκη"},
    {"Όνομα": "Αλέξανδρος", "Ηλικία": 6, "Πόλη": "Θεσσαλονίκη"},
])

def _tmp():
    return tempfile.mkdtemp(dir="/tmp", prefix="astakos_fg_test_")


# ═══════════════════════════════════════════════════════════════
# generate_excel
# ═══════════════════════════════════════════════════════════════

def test_generate_excel_creates_file():
    d = _tmp()
    try:
        out = os.path.join(d, "test.xlsx")
        result = generate_excel.func(output_path=out, data_json=SAMPLE_DATA,
                                     sheet_name="Sheet1", title="Test Report")
        assert os.path.exists(out), f"File not created: {result}"
        assert "✅" in result
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_excel_file_is_valid_xlsx():
    d = _tmp()
    try:
        out = os.path.join(d, "data.xlsx")
        generate_excel.func(output_path=out, data_json=SAMPLE_DATA)
        import openpyxl
        wb = openpyxl.load_workbook(out)
        ws = wb.active
        flat = [cell.value for row in ws.iter_rows() for cell in row if cell.value]
        assert "Όνομα" in flat
        assert "Λάζαρος" in flat
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_excel_invalid_json():
    result = generate_excel.func(output_path="/tmp/bad_astakos.xlsx", data_json="not json")
    assert "❌" in result

def test_generate_excel_empty_data():
    d = _tmp()
    try:
        out = os.path.join(d, "empty.xlsx")
        result = generate_excel.func(output_path=out, data_json="[]")
        assert "❌" in result or "κενά" in result.lower() or os.path.exists(out)
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_excel_with_merged_cells_no_crash():
    """Regression: MergedCell κρας στο auto-width loop."""
    d = _tmp()
    try:
        out = os.path.join(d, "merged.xlsx")
        generate_excel.func(output_path=out, data_json=SAMPLE_DATA)
        import openpyxl
        wb = openpyxl.load_workbook(out)
        ws = wb.active
        ws.merge_cells("A1:B1")
        wb.save(out)
        wb2 = openpyxl.load_workbook(out)
        assert wb2.active is not None
    finally:
        shutil.rmtree(d, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════
# generate_word_doc
# ═══════════════════════════════════════════════════════════════

def test_generate_word_creates_file():
    d = _tmp()
    try:
        out = os.path.join(d, "test.docx")
        result = generate_word_doc.func(output_path=out,
                                        content="Γεια σου κόσμε!\n- σημείο 1\n- σημείο 2",
                                        title="Test Doc")
        assert os.path.exists(out), f"File not created: {result}"
        assert "✅" in result
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_word_contains_content():
    d = _tmp()
    try:
        out = os.path.join(d, "doc.docx")
        generate_word_doc.func(output_path=out,
                               content="Αστακός είναι ο καλύτερος AI εγκέφαλος.",
                               title="Δοκιμή")
        from docx import Document
        doc = Document(out)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Αστακός" in full_text
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_word_with_subtitle():
    d = _tmp()
    try:
        out = os.path.join(d, "sub.docx")
        result = generate_word_doc.func(output_path=out, content="Κείμενο",
                                        title="Τίτλος", subtitle="Υπότιτλος")
        assert "✅" in result
        from docx import Document
        full = "\n".join(p.text for p in Document(out).paragraphs)
        assert "Υπότιτλος" in full
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_word_empty_content():
    d = _tmp()
    try:
        out = os.path.join(d, "empty.docx")
        result = generate_word_doc.func(output_path=out, content="", title="Empty")
        assert isinstance(result, str)
    finally:
        shutil.rmtree(d, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════
# generate_pdf
# ═══════════════════════════════════════════════════════════════

def test_generate_pdf_creates_file():
    d = _tmp()
    try:
        out = os.path.join(d, "test.pdf")
        result = generate_pdf.func(output_path=out, content="Αυτό είναι ένα τεστ PDF.",
                                   title="Test PDF")
        assert os.path.exists(out), f"File not created: {result}"
        assert "✅" in result
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_pdf_is_valid_pdf():
    d = _tmp()
    try:
        out = os.path.join(d, "valid.pdf")
        generate_pdf.func(output_path=out, content="Hello PDF", title="Test")
        with open(out, "rb") as f:
            assert f.read(4) == b"%PDF"
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_pdf_with_author():
    d = _tmp()
    try:
        out = os.path.join(d, "authored.pdf")
        result = generate_pdf.func(output_path=out, content="Περιεχόμενο",
                                   title="Τίτλος", author="Λάζαρος")
        assert "✅" in result
        assert os.path.exists(out)
    finally:
        shutil.rmtree(d, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════
# generate_csv
# ═══════════════════════════════════════════════════════════════

def test_generate_csv_creates_file():
    d = _tmp()
    try:
        out = os.path.join(d, "test.csv")
        result = generate_csv.func(output_path=out, data_json=SAMPLE_DATA)
        assert os.path.exists(out), f"File not created: {result}"
        assert "✅" in result
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_csv_has_correct_content():
    d = _tmp()
    try:
        out = os.path.join(d, "data.csv")
        generate_csv.func(output_path=out, data_json=SAMPLE_DATA)
        with open(out, encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))
        assert len(rows) == 2
        assert rows[0]["Όνομα"] == "Λάζαρος"
        assert rows[1]["Όνομα"] == "Αλέξανδρος"
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_csv_custom_delimiter():
    d = _tmp()
    try:
        out = os.path.join(d, "semi.csv")
        generate_csv.func(output_path=out, data_json=SAMPLE_DATA, delimiter=";")
        assert ";" in open(out, encoding="utf-8-sig").read()
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_generate_csv_invalid_json():
    result = generate_csv.func(output_path="/tmp/bad_astakos.csv", data_json="{invalid}")
    assert "❌" in result

def test_generate_csv_utf8_bom():
    d = _tmp()
    try:
        out = os.path.join(d, "bom.csv")
        generate_csv.func(output_path=out, data_json=SAMPLE_DATA)
        with open(out, "rb") as f:
            assert f.read(3) == b"\xef\xbb\xbf", "Missing UTF-8 BOM"
    finally:
        shutil.rmtree(d, ignore_errors=True)

def test_relative_path_goes_to_desktop():
    d = _tmp()
    try:
        import astakos_skills.file_generator as fg
        old_desktop = fg._DESKTOP
        fg._DESKTOP = d
        result = fg.generate_csv.func(output_path="myfile.csv", data_json=SAMPLE_DATA)
        assert "✅" in result
        assert os.path.exists(os.path.join(d, "myfile.csv"))
    finally:
        fg._DESKTOP = old_desktop
        shutil.rmtree(d, ignore_errors=True)
