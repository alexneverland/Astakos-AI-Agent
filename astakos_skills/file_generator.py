# ================================================================
# Project: Astakos AI Agent 🦞
# Module:  File Generator — Excel, Word, PDF, CSV
# Allows Lobster to create files of various types.
#
# Tools:
#   generate_excel    — .xlsx with headers + rows (WARNING)
#   generate_word_doc — .docx with title and content (WARNING)
#   generate_pdf      — .pdf with title and content (WARNING)
#   generate_csv      — .csv from JSON data (WARNING)
# ================================================================

import os
import json
import csv as _csv_module
from datetime import datetime

from langchain_core.tools import tool


# ── Default output folder ────────────────────────────────────────
# Uses the user's Desktop if a full path is not provided.
_DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")


def _resolve_path(output_path: str, default_ext: str) -> str:
    """If output_path is only a filename (without a directory), it places it on the Desktop."""
    if not os.path.isabs(output_path):
        output_path = os.path.join(_DESKTOP, output_path)
    if not output_path.lower().endswith(default_ext):
        output_path += default_ext
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    return output_path


# ── Excel ────────────────────────────────────────────────────────

@tool
def generate_excel(output_path: str, data_json: str, sheet_name: str = "Sheet1", title: str = "") -> str:
    """
    Creates an Excel file (.xlsx) from data.

    Args:
        output_path: Full file path, e.g., 'C:\\Users\\PC\\Desktop\\report.xlsx'.
                     If only a name is provided, it is saved to the Desktop.
        data_json:   JSON string with a list of dicts.
                     Example: '[{"Name":"John","Age":30},{"Name":"Maria","Age":25}]'
        sheet_name:  Sheet name (default: Sheet1).
        title:       Optional title displayed at the top.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return "❌ Λείπει η βιβλιοθήκη openpyxl. Τρέξε: pip install openpyxl"

    try:
        data = json.loads(data_json)
    except json.JSONDecodeError as e:
        return f"❌ Μη έγκυρο JSON: {e}"

    if not data:
        return "❌ Τα δεδομένα είναι κενά."

    output_path = _resolve_path(output_path, ".xlsx")

    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        row_offset = 1

        # Title (optional)
        if title:
            headers = list(data[0].keys())
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
            title_cell = ws.cell(row=1, column=1, value=title)
            title_cell.font = Font(bold=True, size=14, color="FFFFFF")
            title_cell.fill = PatternFill(start_color="1A3A5C", end_color="1A3A5C", fill_type="solid")
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[1].height = 28
            row_offset = 2

        headers = list(data[0].keys())

        # Header row
        header_font  = Font(bold=True, color="FFFFFF")
        header_fill  = PatternFill(start_color="2E4057", end_color="2E4057", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=row_offset, column=col_idx, value=header)
            cell.font  = header_font
            cell.fill  = header_fill
            cell.alignment = header_align
            cell.border = border
        ws.row_dimensions[row_offset].height = 20

        # Data rows
        even_fill = PatternFill(start_color="F0F4F8", end_color="F0F4F8", fill_type="solid")
        for row_idx, row_data in enumerate(data, row_offset + 1):
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=row_data.get(header, ""))
                cell.border = border
                cell.alignment = Alignment(vertical="center")
                if (row_idx % 2) == 0:
                    cell.fill = even_fill

        # Auto column width (skip MergedCell objects)
        from openpyxl.cell.cell import MergedCell
        for col in ws.columns:
            col_letter = None
            max_len = 0
            for cell in col:
                if isinstance(cell, MergedCell):
                    continue
                if col_letter is None:
                    col_letter = cell.column_letter
                if cell.value is not None:
                    max_len = max(max_len, len(str(cell.value)))
            if col_letter:
                ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

        # Freeze header
        ws.freeze_panes = ws.cell(row=row_offset + 1, column=1)

        wb.save(output_path)
        return f"✅ Excel δημιουργήθηκε: {output_path}\n📊 {len(data)} γραμμές · {len(headers)} στήλες"
    except Exception as e:
        return f"❌ Σφάλμα δημιουργίας Excel: {e}"


# ── Word ─────────────────────────────────────────────────────────

@tool
def generate_word_doc(output_path: str, content: str, title: str = "", subtitle: str = "") -> str:
    """
    [WARNING: DO NOT USE. For .docx documents, use the run_officecli tool.]
    Creates a Word file (.docx) from text.

    Args:
        output_path: Full path, e.g., 'C:\\Users\\PC\\Desktop\\report.docx'.
        content:     Document text. Supports:
                     - Empty line = new paragraph
                     - Lines starting with '## ' = Heading 2
                     - Lines starting with '# ' = Heading 1
                     - Lines starting with '- ' or '* ' = bullet
        title:       Document title (at the top, bold large).
        subtitle:    Optional subtitle.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "❌ Λείπει η βιβλιοθήκη python-docx. Τρέξε: pip install python-docx"

    output_path = _resolve_path(output_path, ".docx")

    try:
        doc = Document()

        # Title
        if title:
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True

        if title or subtitle:
            doc.add_paragraph("")  # spacer

        # Parse content
        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(stripped)

        doc.save(output_path)
        lines = len(content.split("\n"))
        return f"✅ Word αρχείο δημιουργήθηκε: {output_path}\n📄 ~{lines} γραμμές"
    except Exception as e:
        return f"❌ Σφάλμα δημιουργίας Word: {e}"


# ── PDF ──────────────────────────────────────────────────────────

@tool
def generate_pdf(output_path: str, content: str, title: str = "", author: str = "Αστακός") -> str:
    """
    Creates a PDF file from text.

    Args:
        output_path: Full path, e.g., 'C:\\Users\\PC\\Desktop\\report.pdf'.
        content:     Text. Supports:
                     - Lines starting with '## ' = section header
                     - Lines starting with '- ' = bullet
                     - Empty lines = spacer
        title:       Title (at the top).
        author:      Metadata author (default: Lobster).
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "❌ Λείπει η βιβλιοθήκη reportlab. Τρέξε: pip install reportlab"

    output_path = _resolve_path(output_path, ".pdf")

    try:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2.5*cm, bottomMargin=2*cm,
            title=title or "Αστακός Document",
            author=author,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "AstakosTitle",
            parent=styles["Title"],
            fontSize=20,
            textColor=colors.HexColor("#1A3A5C"),
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            "AstakosHeading",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=colors.HexColor("#2E4057"),
            spaceBefore=12,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "AstakosBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "AstakosBullet",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            leftIndent=16,
            bulletIndent=4,
            spaceAfter=4,
        )

        story = []

        if title:
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3*cm))

        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], heading_style))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], heading_style))
            elif stripped.startswith(("- ", "* ")):
                story.append(Paragraph(f"• {stripped[2:]}", bullet_style))
            elif stripped == "":
                story.append(Spacer(1, 0.25*cm))
            else:
                # Escape HTML-like chars for reportlab
                safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))

        doc.build(story)
        return f"✅ PDF δημιουργήθηκε: {output_path}"
    except Exception as e:
        return f"❌ Σφάλμα δημιουργίας PDF: {e}"


# ── CSV ──────────────────────────────────────────────────────────

@tool
def generate_csv(output_path: str, data_json: str, delimiter: str = ",") -> str:
    """
    Creates a CSV file from data.

    Args:
        output_path: Full path, e.g. 'C:\\Users\\PC\\Desktop\\data.csv'.
        data_json:   JSON string with a list of dicts.
                     Example: '[{"Code":"001","Price":9.99}]'
        delimiter:   Delimiter (default: ',').
    """
    try:
        data = json.loads(data_json)
    except json.JSONDecodeError as e:
        return f"❌ Μη έγκυρο JSON: {e}"

    if not data:
        return "❌ Τα δεδομένα είναι κενά."

    output_path = _resolve_path(output_path, ".csv")

    try:
        headers = list(data[0].keys())
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = _csv_module.DictWriter(f, fieldnames=headers, delimiter=delimiter)
            writer.writeheader()
            writer.writerows(data)
        return f"✅ CSV δημιουργήθηκε: {output_path}\n📋 {len(data)} γραμμές · {len(headers)} στήλες"
    except Exception as e:
        return f"❌ Σφάλμα δημιουργίας CSV: {e}"
