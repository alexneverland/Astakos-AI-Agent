# ================================================================
# Project: Astakos AI Agent 🦞
# Developer: User
# Description: Modular LLM-agnostic multi-agent framework
# Copyright (c) 2026 - All Rights Reserved
# ================================================================

import os
import re
import json
import ast
import tempfile
from core.i18n import t
import sys
import math
import subprocess
import base64
import unicodedata
from types import SimpleNamespace
from datetime import datetime, timedelta
from email.message import EmailMessage
from config import NLP_CONFIG
from services.routine_intent import classify_routine_intent
from langchain_core.tools import tool
from pypdf import PdfReader
from google.oauth2.credentials import Credentials
from google.auth.exceptions import RefreshError
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
import spotipy
from spotipy.oauth2 import SpotifyOAuth
import docx
import pandas as pd
import sqlite3
from config import (
    STATE_DB, WORKSPACE_DIR, PHOTOS_INDEX_FILE, PHOTOS_DIR,
    EMAIL_ADDRESS, EMAIL_PASSWORD, GITHUB_TOKEN, VACUUM_IP, VACUUM_TOKEN, GPS_STORAGE_FILE
)
from astakos_skills.linkedin_state_manager import update_pending_linkedin_post, process_and_clear_linkedin_post
from astakos_skills.research_last30days import research_last30days
from memory.vector_store import vector_store, vector_lock, memory, delete_profile_facts_by_exact_fact
_lexical_cache: dict = {}  # {cache_key: (timestamp, data)} — TTL 60s
from services.embeddings import embeddings
from tools.web import (
    get_news, get_weather_forecast, search_supermarket_prices,
    search_goldmall_offers, execute_local_pipeline, get_navigation_info,
    relay_local_payload, search_google_places, browse_url, duckduckgo_search
)
from astakos_skills.morning_briefing import morning_briefing
from astakos_skills.hn_briefing import hn_briefing
from astakos_skills.search_flights import search_flights
from astakos_skills.recipe_expert import recipe_expert, log_meal
from astakos_skills.recipe_library import (
    search_recipe_library,
    get_saved_recipe,
    mark_recipe_favorite,
)
from astakos_skills.repo_mapper import repo_mapper
from tools.project_tools import (
    grant_project_access, list_project_files, read_project_file,
    edit_project_file, write_project_file, grep_project_files,
    list_recent_files,
)
from astakos_skills.file_generator import (
    generate_excel, generate_word_doc, generate_pdf, generate_csv,
)
from astakos_skills.register_tool import register_tool
from astakos_skills.get_world_time import get_world_time
from astakos_skills.manage_context_flag import manage_context_flag
from astakos_skills.text_stats import text_stats
from astakos_skills.scan_receipt import scan_receipt
from astakos_skills.officecli_skill import run_officecli
from astakos_skills.read_agent_skill import list_agent_skills, read_agent_skill

# ────────────────────────────────────────────────────────────────
# CREDENTIALS PATHS
# ────────────────────────────────────────────────────────────────
import config
TOKEN_PATH = config.TOKEN_PATH
CREDS_PATH = config.CREDENTIALS_PATH

# ────────────────────────────────────────────────────────────────
# PROTECTED SANDBOX
# ────────────────────────────────────────────────────────────────
PROTECTED_FILES = ["main.py", "telegram_bot.py", "update.py", ".env"]
DANGEROUS_WORDS = [
    "os.remove", "os.rmdir", "shutil.rmtree", "format c:",
    "exec(", "eval(", "compile(", "__import__", "subprocess.run",
    "subprocess.call", "subprocess.Popen", "os.system"
]


# ────────────────────────────────────────────────────────────────
# MEMORY TOOLS
# ────────────────────────────────────────────────────────────────
@tool
def archive_file(filename: str, content_summary: str) -> str:
    """
    Permanently archives a file (photo, document, PDF) in memory (JSON + ChromaDB).
    filename: The exact technical name of the file (e.g., web_xxx.pdf or web_xxx.png).
    content_summary: Summary of the document's content or the analysis of the image.
    """
    try:
        import os
        from config import BASE_DIR, PHOTOS_DIR
        from memory.vector_store import memory

        search_dirs = [
            PHOTOS_DIR,
            os.path.join(BASE_DIR, "outputs"),
            os.path.join(BASE_DIR, "telegram_uploads"),
            os.path.join(BASE_DIR, "telegram_photos"),
            os.path.join(BASE_DIR, "uploads")
        ]

        full_path = None
        for d in search_dirs:
            test_path = os.path.join(d, filename)
            if os.path.exists(test_path) and os.path.isfile(test_path):
                full_path = test_path
                break

        if not full_path:
            return t("tools.system.archive_not_found", filename=filename)

        ext = os.path.splitext(full_path)[1].lower()
        m_type = "photo" if ext in [".jpg", ".jpeg", ".png", ".webp", ".gif"] else "document"

        memory.save(
            memory_type=m_type,
            file_path=full_path,
            analysis=content_summary,
            caption=f"Archive ({m_type}): {filename}"
        )
        return t("tools.system.archive_success", filename=filename, m_type=m_type)
    except Exception as e:
        return t("tools.system.archive_error", error=str(e))

# Channel for Memory Provenance — defined by server.py/telegram_bot.py
_CURRENT_CHANNEL: str = "unknown"


def _normalize_memory_query(value: str) -> str:
    text = unicodedata.normalize("NFD", str(value or "").lower())
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return " ".join(text.split())


def _expand_memory_query(query: str) -> tuple[list[str], str | None]:
    clean = _normalize_memory_query(query)
    expanded = [query]

    family_markers = t("tools.system.family_markers")
    project_markers = t("tools.system.project_markers")
    home_markers = t("tools.system.home_markers")
    lesson_markers = t("tools.system.lesson_markers")

    inferred_category = None
    if any(marker in clean for marker in family_markers):
        inferred_category = "family"
        expanded.append(t("tools.system.family_expanded", query=query))
    elif any(marker in clean for marker in project_markers):
        inferred_category = "projects"
        expanded.append(t("tools.system.project_expanded", query=query))
    elif any(marker in clean for marker in home_markers):
        inferred_category = "home"
        expanded.append(t("tools.system.home_expanded", query=query))
    elif any(marker in clean for marker in lesson_markers):
        inferred_category = "lesson"
        expanded.append(t("tools.system.lesson_expanded", query=query))

    is_gift_or_product = any(
        marker in clean
        for marker in t("tools.system.gift_markers")
    )
    if is_gift_or_product:
        expanded.append(t("tools.system.gift_expanded", query=query))
        if inferred_category is None and any(marker in clean for marker in family_markers):
            inferred_category = "family"

    # Keep order, remove exact duplicates.
    unique = []
    seen = set()
    for item in expanded:
        key = _normalize_memory_query(item)
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    return unique, inferred_category


def _memory_query_tokens(query: str) -> list[str]:
    stopwords = set(t("tools.system.stopwords"))
    tokens = re.findall(t("tools.system.greek_words_regex"), _normalize_memory_query(query))
    return [token for token in tokens if len(token) >= 4 and token not in stopwords]


def _stem_token(token: str) -> str:
    """Rough Greek stemming: cuts off the most common inflectional endings
    (cases/number) so that suffixes match.
    Always keeps a stem of >= 4 characters (the same limit as tokens) to
    prevent noise from increasing due to very short stems.
    """
    if len(token) >= 7:
        return token[:-2]
    if len(token) >= 5:
        return token[:-1]
    return token


def _lexical_memory_matches(query: str, category: str = "", limit: int = 4) -> list:
    """Keyword fallback over Chroma docs; complements embeddings for exact user terms.
    L1 cache (60s TTL) avoids full collection.get() on every call."""
    import time as _time
    tokens = _memory_query_tokens(query)
    if len(tokens) < 2:
        return []
    try:
        cache_key = category or "__all__"
        cached = _lexical_cache.get(cache_key)
        if cached and (_time.monotonic() - cached[0]) < 60:
            data = cached[1]
        else:
            kwargs = {"include": ["documents", "metadatas"]}
            if category:
                kwargs["where"] = {"category": category}
            data = vector_store._collection.get(**kwargs)
            _lexical_cache[cache_key] = (_time.monotonic(), data)
    except Exception:
        return []

    clean_query = _normalize_memory_query(query)
    scored = []
    for document, metadata in zip(data.get("documents", []), data.get("metadatas", [])):
        clean_doc = _normalize_memory_query(document)
        score = sum(1 for token in tokens if _stem_token(token) in clean_doc)
        if any(m in clean_query for m in t("tools.system.word_link_markers")):
            score += 1 if ("http" in clean_doc or "link" in clean_doc) else 0
        if t("tools.system.word_gift") in clean_query:
            score += 1 if (t("tools.system.word_gift") in clean_doc or t("tools.system.word_buy") in clean_doc or t("tools.system.word_future") in clean_doc) else 0
        if t("tools.system.word_birthday") in clean_query:
            score += 1 if (t("tools.system.word_birthday") in clean_doc or t("tools.system.word_reminder") in clean_doc) else 0
        if score < 2:
            continue
        scored.append((score, len(str(document)), document, metadata or {}))

    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [
        SimpleNamespace(page_content=document, metadata=metadata)
        for _, _, document, metadata in scored[:limit]
    ]

@tool
def search_memory(query: str, category: str = "") -> str:
    """Search in long-term memory. Call this ONCE ONLY. If you already have [Information from search] in the context DO NOT call it again. Use it before answering:
    1. Questions about {config.USER_NAME}, family, home, habits, or projects.
    2. Issues that require suggestions, advice, or solutions.
    3. References to the past or to existing equipment.

    Args:
        query: Keywords (e.g., 'kid food', 'project backend')
        category: Optional filter: 'lazaros', 'family', 'projects', 'home', 'lesson', 'photos'
    """
    VALID_CATS = {"lazaros", "family", "projects", "home", "lesson", "session", "photos"}
    try:
        search_queries, inferred_category = _expand_memory_query(query)
        primary_query = search_queries[-1]
        effective_category = category if category in VALID_CATS else (inferred_category or "")
        sql_lines = []
        try:
            from memory.context_builder import temporal_history_for_query
            from tools import system as _self

            sql_lines = temporal_history_for_query(
                primary_query,
                channel=getattr(_self, "_CURRENT_CHANNEL", "telegram") or "telegram",
                limit=8,
            )
        except Exception:
            sql_lines = []

        with vector_lock:
            merged_results = []
            seen_docs = set()
            for doc in _lexical_memory_matches(primary_query, effective_category, limit=4):
                key = getattr(doc, "page_content", str(doc))
                if key in seen_docs:
                    continue
                seen_docs.add(key)
                merged_results.append(doc)
            # [PERF]: 1 similarity_search instead of 3 — primary_query is sufficient (expanded queries do not improve significantly)
            for search_query in search_queries[:1]:
                if effective_category:
                    batch = vector_store.similarity_search(search_query, k=6, filter={"category": effective_category})
                else:
                    batch = vector_store.similarity_search(search_query, k=6)
                for doc in batch:
                    key = getattr(doc, "page_content", str(doc))
                    if key in seen_docs:
                        continue
                    seen_docs.add(key)
                    merged_results.append(doc)
            results = merged_results[:8 if effective_category else 6]

        from memory.vector_store import (
            get_latest_state_for_query,
            build_profile_memory_summary,
        )
        latest = get_latest_state_for_query(primary_query, category=effective_category or None)
        profile_lines = build_profile_memory_summary(
            primary_query,
            category=effective_category or None,
            limit=5,
        )

        if not results and not sql_lines and not latest and not profile_lines:
            return t("tools.system.memory_no_results")

        # bump retrieval_count async — does not block the response
        if results:
            import threading as _thr
            def _bump_async():
                try:
                    from memory.vector_store import bump_retrieval_count
                    with vector_lock:
                        kwargs = {"n_results": min(6, len(results))}
                        if effective_category:
                            kwargs["where"] = {"category": effective_category}
                        raw = vector_store._collection.query(
                            query_embeddings=[embeddings.embed_query(primary_query)], **kwargs
                        )
                    if raw.get("ids") and raw["ids"][0]:
                        bump_retrieval_count(raw["ids"][0])
                except Exception:
                    pass
            _thr.Thread(target=_bump_async, daemon=True).start()

        by_cat: dict = {}
        for res in results:
            cat = res.metadata.get("category", "general")
            content = res.page_content
            photo_path = res.metadata.get("photo_path")
            if photo_path:
                content += f" [PHOTO_PATH: {photo_path}]"
            by_cat.setdefault(cat, []).append(content)

        output_parts = [t("tools.system.memory_found")]
        
        # 1. Profile: Latest matching state (generic, query-driven)
        if latest and latest.get("fact"):
            output_parts.append("\n[LATEST MATCHING STATE]")
            output_parts.append(f"  • {latest['fact']}")

        # 2. Structured profile memory summary
        if profile_lines:
            output_parts.append("\n[STRUCTURED PROFILE MEMORY]")
            output_parts.extend(profile_lines)

        if sql_lines:
            output_parts.append(t("tools.system.sqlite_history"))
            output_parts.extend(f"  • {line}" for line in sql_lines)

        output_parts.append(t("tools.system.chroma_memory"))
        if by_cat:
            for cat, facts in by_cat.items():
                output_parts.append(f"\n[{cat.upper()}]")
                output_parts.extend(f"  • {f}" for f in facts)
        else:
            output_parts.append(t("tools.system.no_chroma"))

        output_parts.append(
            t("tools.system.instruction")
        )
        return "\n".join(output_parts).strip()
    except Exception as e:
        return t("tools.system.memory_error", error=str(e))
@tool
def run_terminal_command(command: str, already_approved: bool = False) -> str:
    """
    Executes PowerShell commands on {config.USER_NAME}'s PC ({config.DEVELOPER_NAME}) and returns the result.
    Ideal for:
    - Reading logs (e.g., 'Get-Content C:\\path\\to\\mastroapp\\logs\\error.log -Tail 50').
    - Checking ports (e.g., 'netstat -ano | findstr 8000').
    - Running tests (e.g., 'python manage.py test').
    already_approved=True: bypasses only the confirmation gate, not BLOCKED commands.
    """
    import subprocess
    from core.safe_executor import safe_execute, classify_command, ExecPolicy
    print(f"\033[93m[Terminal Execution]: {command}\033[0m")

    def _executor(cmd):
        try:
            result = subprocess.run(
                ["powershell", "-Command", cmd],
                capture_output=True, text=True, timeout=30,
                encoding='utf-8', errors='ignore'
            )
            output = result.stdout or ""
            if result.returncode != 0:
                output += f"\nERROR:\n{result.stderr}"
            elif result.stderr:
                output += f"\nWARNINGS:\n{result.stderr}"

            if not output.strip():
                return {"status": "ok", "output": "[SUCCESS] Command executed successfully (no output). PLEASE PROCEED TO THE NEXT STEP/TOOL."}

            if len(output) > 10000:
                output = output[:10000] + "\n... [output truncated]"

            return {"status": "ok", "output": f"💻 Terminal Output:\n{output}"}
        except subprocess.TimeoutExpired:
            return {"status": "ok", "output": "❌ Timeout: >30 seconds."}
        except Exception as e:
            return {"status": "ok", "output": f"❌ Terminal Error: {str(e)}"}

    # Even after approval, hard-blocked commands are never executed.
    if already_approved:
        policy, reason = classify_command(command)
        if policy == ExecPolicy.BLOCKED:
            result = {"status": "blocked", "reason": reason}
        else:
            result = _executor(command)
    else:
        result = safe_execute(command, _executor)

    if result.get("status") == "blocked":
        return f"🛡️ [SAFE EXECUTOR - BLOCKED]: {result['reason']}"
    if result.get("status") == "cancelled":
        return f"⚠️ [SAFE EXECUTOR]: Command requires confirmation. Send again with `/confirm {command}`"
    return result.get("output", "")

@tool
def save_to_memory(
    fact: str,
    entities: str = "",
    category: str = "other",
    reason: str = "agent_inferred",
    external_content_sources_json: str = "",
) -> str:
    """
    Saves information SEMANTICALLY.
    fact: The fact (e.g., "Kid1 only eats lentils").
    entities: Keywords separated by commas (e.g., "Kid1, Food, Preference").
    category: The category (e.g., 'family', 'home', 'lazaros', 'tech', 'work').
    reason: Why it is being saved — 'user_stated' if explicitly said by the user, 'agent_inferred' otherwise.
    external_content_sources_json: Internal approval provenance. Do not set this manually.

    ⚡ Fire-and-forget: ChromaDB/Vertex AI work is done in a background thread.
    Returns immediately so that the agent does not block the user for ~11s.
    """
    import datetime
    import threading
    from tools import system as _self

    # Capture channel context BEFORE the thread (module-level var, thread-unsafe if read later)
    _source = _self._CURRENT_CHANNEL

    def _do_save():
        try:
            from memory.vector_store import memory
            from memory.session_memory import build_canonical_memory_candidate

            canonical_fact = fact.strip()
            if not canonical_fact.startswith(("[USER_FACT]", "[LESSON]", "[CAPABILITY]")):
                canonical_fact = f"[USER_FACT]: {canonical_fact}"

            raw_entities = [x.strip() for x in entities.split(",") if x.strip()]
            from core.untrusted_content import external_content_sources_from_json

            external_content_sources = external_content_sources_from_json(
                external_content_sources_json,
            )

            candidate = build_canonical_memory_candidate(
                memory_type="fact",
                fact=canonical_fact,
                category=category,
                entities=raw_entities,
                agent_name="Tool_save_to_memory",
                source=_source,
                reason=reason,
                confidence=0.85 if reason == "user_stated" else 0.7,
            )
            if external_content_sources:
                candidate["external_content_sources"] = external_content_sources

            saved = memory.save(**candidate)

            if saved:
                _lexical_cache.clear()
        except Exception as e:
            print(f"⚠️ [save_to_memory bg]: {e}")

    threading.Thread(target=_do_save, daemon=True).start()
    return f"✅ Saving in background: [{entities}]"


@tool
def delete_from_memory(query: str) -> str:
    """PERMANENTLY deletes information from memory (Chroma).

    USE THIS TOOL (not save_to_memory) whenever the user explicitly requests
    to erase/delete/remove already stored information because it is
    incorrect, outdated, or irrelevant — e.g., "erase the memory that says X", "delete this
    about Y", "what I said about X is wrong, remove it" etc.
    The save_to_memory tool ONLY adds new entries; upon a deletion/correction request,
    it leaves the incorrect entry in place — which is why in such cases you should always
    prefer delete_from_memory (if needed, you can subsequently call
    save_to_memory with the correct content in a separate step).

    query: Provide a BRIEF, SPECIFIC phrase that identifies ONLY the incorrect entry
    (e.g., "old wrong address"), not the user's entire sentence/correction.
    """
    try:
        def _norm(t: str) -> str:
            t = unicodedata.normalize("NFD", str(t or "").lower())
            return "".join(ch for ch in t if unicodedata.category(ch) != "Mn")

        norm_query = _norm(query).strip()

        with vector_lock:
            collection = vector_store._collection
            data = collection.get(include=["documents", "metadatas"])

        # 1) Exact substring match FIRST — more reliable than embeddings when_
        # the phrases are close/similar (e.g. old incorrect vs correct address:
        # the embeddings see them as almost identical and a wrong record might be deleted).
        literal_hits = [
            (doc_id, doc) for doc_id, doc in zip(data.get("ids", []), data.get("documents", []))
            if norm_query and norm_query in _norm(doc)
        ]

        if len(literal_hits) == 1:
            target_id, content = literal_hits[0]
            with vector_lock:
                collection.delete(ids=[target_id])
            profile_deleted = delete_profile_facts_by_exact_fact(content)
            print(f"\n🔥 [DATABASE ACTION]: DELETED (exact match): {content}")
            return (
                f"Memory '{content}' deleted successfully "
                f"(Chroma + {profile_deleted} structured profile record(s))."
            )

        if len(literal_hits) > 1:
            previews = "\n".join(f"  • {str(c).strip()[:140]}" for _, c in literal_hits[:6])
            return (
                f"⚠️ Found {len(literal_hits)} records matching '{query}'. "
                f"Be more specific about which to delete:\n{previews}"
            )

        # 2) Fallback: semantic search (embeddings), only when no
        # there is no literal match. u_00ad_ u_00ad__
        query_emb = embeddings.embed_query(query)
        with vector_lock:
            results = collection.query(query_embeddings=[query_emb], n_results=1)

            if not results['ids'] or not results['ids'][0]:
                return "No relevant records found for deletion."

            content = results['documents'][0][0]
            distance = results['distances'][0][0] if 'distances' in results and results['distances'] else 1.0

            if distance > 0.40:
                return (
                    f"⚠️ Not deleted. Closest match (Distance: {distance:.2f}): "
                    f"'{content}'. Please be more specific."
                )

            target_id = results['ids'][0][0]
            collection.delete(ids=[target_id])
            profile_deleted = delete_profile_facts_by_exact_fact(content)

        print(f"\n🔥 [DATABASE ACTION]: DELETED (Dist: {distance:.2f}): {content}")
        return (
            f"Memory '{content}' deleted successfully "
            f"(Chroma + {profile_deleted} structured profile record(s))."
        )
    except Exception as e:
        return f"Deletion error: {e}"


@tool
def retrieve_photo(query: str) -> str:
    """Retrieves a photo from memory. WHEN it returns [SEND_PHOTO: path], INCLUDE IT EXACTLY AS IS in your response."""
    try:
        import numpy as np

        with vector_lock:
            results = vector_store.similarity_search(query, k=10)

        for doc in results:
            photo_path = doc.metadata.get("photo_path")
            if photo_path and os.path.exists(photo_path):
                return (
                    f"Found the photo!\n"
                    f"Description: {doc.page_content}\n"
                    f"[SEND_PHOTO: {photo_path}]"
                )

        if os.path.exists(PHOTOS_INDEX_FILE):
            with open(PHOTOS_INDEX_FILE, "r", encoding="utf-8") as f:
                index = json.load(f)

            if index:
                query_emb = np.array(embeddings.embed_query(query))
                best_score = -1.0
                best_entry = None

                for entry in index:
                    candidate = f"{entry.get('caption', '')} {entry.get('analysis', '')}".strip()
                    if not candidate:
                        continue
                    cand_emb = np.array(embeddings.embed_query(candidate))
                    norm_q = np.linalg.norm(query_emb)
                    norm_c = np.linalg.norm(cand_emb)
                    if norm_q and norm_c:
                        sim = float(np.dot(query_emb, cand_emb) / (norm_q * norm_c))
                        if sim > best_score:
                            best_score = sim
                            best_entry = entry

                if best_score < 0.35:
                    return "System: No relevant photo found for this query."

                if best_entry:
                    fp = best_entry.get("file_path", "")
                    note = "" if best_score >= 0.5 else " (No exact match found — providing the closest one.)"
                    if not fp:
                        best_entry = index[-1]
                        fp = best_entry.get("file_path", "")
                        note = " (Fallback: most recent photo.)"

                    if fp and os.path.exists(fp):
                        return (
                            f"Found photo from {best_entry.get('date', 'unknown date')}{note}\n"
                            f"[SEND_PHOTO: {fp}]"
                        )

        return "System: Photo not found."

    except Exception as e:
        return f"Error: Failed to retrieve photo: {str(e)}"


# ────────────────────────────────────────────────────────────────
# REMINDERS & LISTS
# ────────────────────────────────────────────────────────────────

def _normalize_reminder_text(text: str) -> str:
    import re

    value = str(text or "").lower().strip()

    replacements = t("tools.system.reminder_prefixes")

    for old, new in replacements.items():
        value = value.replace(old, new)

    value = re.sub(t("tools.system.greek_chars_regex"), " ", value, flags=re.UNICODE)
    value = re.sub(r"\s+", " ", value).strip()
    return value

def _same_reminder_task(existing_task: str, new_task: str) -> bool:
    a = _normalize_reminder_text(existing_task)
    b = _normalize_reminder_text(new_task)

    if not a or not b:
        return False

    if a == b:
        return True

    a_tokens = {tok for tok in a.split() if len(tok) >= 4}
    b_tokens = {tok for tok in b.split() if len(tok) >= 4}

    if not a_tokens or not b_tokens:
        return False

    overlap = len(a_tokens & b_tokens)
    min_len = min(len(a_tokens), len(b_tokens))

    return overlap >= max(2, min_len)


def _same_pending_reminder(existing_task: str, new_task: str, existing_time: str, new_time: str) -> bool:
    return (
        str(existing_time or "").strip() == str(new_time or "").strip()
        and _same_reminder_task(existing_task, new_task)
    )

@tool
def set_local_reminder(
    task: str,
    minutes_from_now: int = 0,
    exact_time: str = None,
    action: str = "add",
    location: str = None,
    match_task: str = None,
    external_content_sources_json: str = "",
) -> str:
    """
    Manages local reminders.
    action: 'add' (new), 'read' (read pending ONLY), 'done' (completion), 'update' (correct text)
    task: For 'add'/'update' → description. For 'done' → keyword of the reminder being completed.
    match_task: For 'update' → the current reminder description to replace. Read first;
                if multiple pending reminders match, ask the user to clarify rather than guessing.
    location: ONLY for location-based reminders. Use 'home' for arrival home,
              or 'leave_current_location' to trigger after leaving the current place.
              When location is provided, DO NOT provide minutes_from_now or exact_time.
    external_content_sources_json: Internal approval provenance. Do not set it manually.
    """
    conn = None
    try:
        conn = sqlite3.connect(STATE_DB)
        cursor = conn.cursor()

        # ── READ: Returns ONLY pending ──────────────────────
        if action == "read":
            cursor.execute(
                "SELECT task, time, external_content_sources_json "
                "FROM reminders WHERE status='pending'"
            )
            pending = cursor.fetchall()
            if not pending:
                return t("tools.system.reminders_read_empty")
            lines = []
            for rtask, tm, sources_json in pending:
                from core.untrusted_content import (
                    external_content_sources_from_json,
                    format_untrusted_tool_result,
                )

                sources = external_content_sources_from_json(sources_json or "")
                if sources:
                    rtask = format_untrusted_tool_result(
                        f"persisted reminder sources: {', '.join(sources)}",
                        rtask,
                    )
                if tm and tm.startswith("loc:"):
                    loc = tm.split(":", 1)[1]
                    if loc == "leave_current_location":
                        loc = t("tools.system.reminders_location_leave_current")
                    lines.append(f"• [📍 {loc}] {rtask}")
                else:
                    lines.append(f"• [{tm}] {rtask}")
            return t("tools.system.reminders_read_header") + "\n".join(lines)

        # ── DONE: Closes reminder with keyword ────────────────
        elif action == "done":
            cursor.execute("SELECT id, task FROM reminders WHERE status='pending'")
            pending = cursor.fetchall()
            found_id = None
            for rid, rtask in pending:
                if task.lower() in rtask.lower():
                    found_id = rid
                    break
            
            if not found_id:
                return t("tools.system.reminders_done_not_found", task=task)
                
            cursor.execute("UPDATE reminders SET status='done' WHERE id=?", (found_id,))
            conn.commit()
            return t("tools.system.reminders_done_success", task=task)

        elif action == "update":
            if not match_task:
                return t("tools.system.reminders_update_req_match")

            cursor.execute("SELECT id, task, time FROM reminders WHERE status='pending'")
            matches = [
                row for row in cursor.fetchall()
                if _same_reminder_task(row[1], match_task)
            ]
            if not matches:
                return t("tools.system.reminders_update_not_found", task=match_task)
            if len(matches) > 1:
                options = "\n".join(f"• [{tm}] {existing_task}" for _, existing_task, tm in matches)
                return t("tools.system.reminders_update_ambiguous", options=options)

            reminder_id, _, reminder_time = matches[0]
            cursor.execute(
                "SELECT task FROM reminders WHERE status='pending' AND id != ? AND time = ?",
                (reminder_id, reminder_time),
            )
            for (existing_task,) in cursor.fetchall():
                if _same_reminder_task(existing_task, task):
                    return t(
                        "tools.system.reminders_add_exists",
                        target_time=reminder_time,
                        existing_task=existing_task,
                    )

            cursor.execute(
                "SELECT external_content_sources_json FROM reminders WHERE id=?",
                (reminder_id,),
            )
            existing_sources_json = cursor.fetchone()[0]
            from core.untrusted_content import external_content_sources_from_json

            existing_sources = external_content_sources_from_json(existing_sources_json or "")
            new_sources = external_content_sources_from_json(external_content_sources_json)
            cursor.execute(
                "UPDATE reminders SET task=?, external_content_sources_json=? WHERE id=?",
                (task, json.dumps(sorted(set(existing_sources) | set(new_sources))), reminder_id),
            )
            conn.commit()
            return t("tools.system.reminders_update_success", task=task)

        # ── ADD: New reminder ─────────────────────────────────
        else:
            from datetime import datetime, timedelta
            from memory.location_reminders import (
                LEAVE_CURRENT_LOCATION,
                get_fresh_current_location,
                save_leave_current_location_anchor,
            )

            current_location = None
            from core.untrusted_content import external_content_sources_from_json

            provenance_json = json.dumps(
                external_content_sources_from_json(external_content_sources_json)
            )

            if minutes_from_now > 0:
                target_time = (datetime.now() + timedelta(minutes=minutes_from_now)).strftime("%Y-%m-%d %H:%M")
            elif exact_time:
                exact_time = exact_time.strip()
                if len(exact_time) <= 5 and ":" in exact_time:
                    today_str = datetime.now().strftime("%Y-%m-%d")
                    target_time = f"{today_str} {exact_time}"
                else:
                    try:
                        datetime.strptime(exact_time, "%Y-%m-%d %H:%M")
                        target_time = exact_time
                    except ValueError:
                        return t("tools.system.reminders_add_err_time")
            elif location:
                if location == LEAVE_CURRENT_LOCATION:
                    current_location = get_fresh_current_location()
                    if current_location is None:
                        return t("tools.system.reminders_add_err_live_location")
                target_time = f"loc:{location}"
            else:
                return t("tools.system.reminders_add_err_args")

            cursor.execute("SELECT id, task, time FROM reminders WHERE status='pending'")
            pending_rows = cursor.fetchall()

            related_reminders = []
            for rid, existing_task, existing_time in pending_rows:
                if _same_pending_reminder(existing_task, task, existing_time, target_time):
                    return t("tools.system.reminders_add_exists", target_time=target_time, existing_task=existing_task)
                if _same_reminder_task(existing_task, task):
                    related_reminders.append((existing_task, existing_time))

            cursor.execute(
                "INSERT INTO reminders (task, time, status, external_content_sources_json) "
                "VALUES (?, ?, 'pending', ?)",
                (task, target_time, provenance_json),
            )
            if current_location is not None:
                save_leave_current_location_anchor(
                    conn,
                    reminder_id=cursor.lastrowid,
                    anchor_lat=current_location[0],
                    anchor_lon=current_location[1],
                )
            conn.commit()

            if location:
                if location == LEAVE_CURRENT_LOCATION:
                    response = t("tools.system.reminders_add_success_leave_current")
                else:
                    response = t("tools.system.reminders_add_success_loc", location=location)
            else:
                response = t("tools.system.reminders_add_success_time", target_time=target_time)

            if related_reminders:
                existing_task, existing_time = related_reminders[0]
                response += "\n" + t(
                    "tools.system.reminders_add_related_existing",
                    existing_task=existing_task,
                    existing_time=existing_time,
                )
            return response

    except Exception as e:
        return t("tools.system.reminders_err_generic", e=str(e))
    finally:
        if conn:
            conn.close()
from langchain_core.tools import tool
from memory.routine_db import upsert_routine

@tool
def learn_routine(
    day_of_week: str,
    time_str: str,
    event_name: str,
    event_type: str = "general",
    external_content_sources_json: str = "",
) -> str:
    """
    [CRITICAL]: Use this WHEN {config.USER_NAME} mentions a habit,
    a routine, or something that is repeated (e.g., "Every Friday at 13:00 I go to the farmers market").

    [OSMANI RULE - SEARCH BEFORE EDIT]:
    ALWAYS call `get_routines` BEFORE using this tool to ensure a similar routine doesn't already exist. If it does, use `edit_routine` instead!

    RULES FOR ARGUMENTS:
    - day_of_week: English canonical ("Monday"…"Sunday") or "Everyday" or "Weekdays".
    - time_str: Time in HH:MM (e.g., "13:00"). If no time is mentioned, DO NOT call the tool.
    - event_name: BRIEF canonical description in 2-4 words (e.g., "message Kostas", "farmers market",
      "gym"). DO NOT include "Every day", "Every morning", or time phrases — these belong
      to day_of_week/time_str. The event_name must be CONSISTENT for the same activity.
    - event_type: "family", "work", "hobby", "general".

    ATTENTION: Call this ONLY for recurring activities. Ignore one-off events
    ("today I went…", "tomorrow I have…").
    """
    from datetime import datetime
    from core.untrusted_content import external_content_sources_from_json

    VALID_DAYS = {"Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Everyday", "Weekdays"}
    VALID_TYPES = {"family", "work", "hobby", "general"}

    if day_of_week not in VALID_DAYS:
        return t("tools.system.routine_err_day", day=day_of_week)

    try:
        datetime.strptime(time_str, "%H:%M")
    except ValueError:
        return t("tools.system.routine_err_time", time_str=time_str)

    if len(event_name.strip()) < 3:
        return t("tools.system.routine_err_name")

    if event_type not in VALID_TYPES:
        event_type = "general"

    try:
        res = upsert_routine(
            day_of_week,
            time_str,
            event_name,
            event_type,
            confidence_boost=0.3,
            external_content_sources=external_content_sources_from_json(
                external_content_sources_json,
            ),
        )

        if res == "created":
            return t("tools.system.routine_added_first", event_name=event_name)
        elif res == "merged":
            return t("tools.system.routine_added_merged", event_name=event_name)
        else:
            return t("tools.system.routine_added_boosted", event_name=event_name)
    except Exception as e:
        return t("tools.system.routine_add_err", e=str(e))


@tool
def delete_routine(event_name: str, day_of_week: str = "", time_str: str = "") -> str:
    """
    [ACTION]: Permanently deletes a routine from the scheduler (routine database).
    Use this when the user explicitly asks to delete, cancel, or abolish something.
    a recurring routine (not a calendar event, not a simple memory, but a routine!).
    - event_name: The name or part of the name of the routine.
    - day_of_week: (Optional) Day for a more precise search.
    - time_str: (Optional) Time (HH:MM) for a more precise search.
    """
    try:
        from memory.routine_db import find_routines_for_schedule_control, delete_routine_db

        routines = find_routines_for_schedule_control(
            event_name, 
            day_of_week=day_of_week if day_of_week else None, 
            time_str=time_str if time_str else None
        )

        if not routines:
            return t("tools.system.routine_not_found_delete", event_name=event_name)

        if len(routines) > 1:
            opts = "\n".join(f"- {r['event']} ({r.get('day','')}, {r.get('time','')})" for r in routines)
            return t("tools.system.routine_multiple_found", opts=opts)

        r_id = routines[0]["id"]
        success = delete_routine_db(r_id)
        if success:
            return t("tools.system.routine_deleted_success", event=routines[0]["event"])
        return t("tools.system.routine_deleted_fail")
    except Exception as e:
        return t("tools.system.routine_delete_err", e=str(e))


@tool
def edit_routine(
    event_name: str,
    new_time_str: str = "",
    new_day_of_week: str = "",
    day_of_week: str = "",
    time_str: str = "",
) -> str:
    """
    [ACTION]: Changes the time and/or day of an existing routine in the scheduler.
    Use this instead of learn_routine when {config.USER_NAME} asks to CHANGE the time 
    or day of something that already exists!
    - event_name: The name (or part) of the existing routine.
    - new_time_str: The new time (e.g., "23:00"). Leave empty if it does not change.
    - new_day_of_week: The new day (e.g., "Everyday", "Weekdays", "Monday"). Leave empty if it does not change.
    - day_of_week: (Optional) Day of the existing routine for clarification.
    - time_str: (Optional) Time of the existing routine for clarification.
    """
    if not new_time_str and not new_day_of_week:
        return t("tools.system.routine_update_req_time_day")

    try:
        from memory.routine_db import find_routines_for_schedule_control, update_routine_db
        import re

        routines = find_routines_for_schedule_control(
            event_name,
            day_of_week=day_of_week if day_of_week else None,
            time_str=time_str if time_str else None,
        )

        if not routines:
            return t("tools.system.routine_not_found", event_name=event_name)

        if len(routines) > 1:
            opts = "\n".join(f"- {r['event']} ({r.get('day','')}, {r.get('time','')})" for r in routines)
            return t("tools.system.routine_multiple_found", opts=opts)

        r_id = routines[0]["id"]
        
        # If new_time_str was provided, make sure it is HH:MM
        if new_time_str:
            if not re.match(r"^([01]\d|2[0-3]):([0-5]\d)$", new_time_str):
                return t("tools.system.routine_err_time", time_str=new_time_str)

        success = update_routine_db(r_id, new_time=new_time_str, new_day=new_day_of_week)
        if success:
            return t("tools.system.routine_update_success", event=routines[0]["event"], new_time=new_time_str or t("tools.system.same"), new_day=new_day_of_week or t("tools.system.same"))
        return t("tools.system.routine_update_fail")
    except Exception as e:
        return t("tools.system.routine_update_err", e=str(e))
@tool
def get_routines(day_of_week: str) -> str:
    """
    [QUERY]: Returns the recorded routines for a specific day.
    Use this when {config.USER_NAME} asks "what do I have on Friday?" or "which routines do you know?".
    - day_of_week: e.g. "Monday", "Friday", "Everyday"
    """
    try:
        from memory.routine_db import get_routines_for_day
        routines = get_routines_for_day(day_of_week)
        if not routines:
            return t("tools.system.routine_none_for_day", day_of_week=day_of_week)
        
        lines = [t("tools.system.routine_day_header", day_of_week=day_of_week)]
        for r in routines:
            conf_pct = int(r['confidence'] * 100)
            mentions = r.get('mentions', 1)
            event = r["event"]
            source_json = r.get("external_content_sources_json", "")
            if source_json:
                from core.untrusted_content import (
                    external_content_sources_from_json,
                    format_untrusted_tool_result,
                )
                sources = external_content_sources_from_json(source_json)
                if sources:
                    event = format_untrusted_tool_result(
                        "persisted routine sources: " + ", ".join(sources),
                        event,
                    )
            lines.append(t("tools.system.routine_day_item", time=r["time"], event=event, type=r["type"], conf=conf_pct, mentions=mentions))
        return "\n".join(lines)
    except Exception as e:
        return t("tools.system.routine_fetch_err", e=str(e))


@tool
def search_routines(event_name: str) -> str:
    """
    [QUERY]: Searches for existing routines across ALL days by keyword or name.
    [OSMANI RULE]: Use this BEFORE calling `learn_routine` to verify if the routine already exists on a different day/time!
    - event_name: The name or part of the name of the routine (e.g. 'Roblox', 'park').
    """
    try:
        from memory.routine_db import find_routines_for_schedule_control
        routines = find_routines_for_schedule_control(event_name)
        if not routines:
            return f"No existing routines found matching '{event_name}'."
        
        lines = [f"Found {len(routines)} matching routines:"]
        for r in routines:
            event = r["event"]
            source_json = r.get("external_content_sources_json", "")
            if source_json:
                from core.untrusted_content import (
                    external_content_sources_from_json,
                    format_untrusted_tool_result,
                )
                sources = external_content_sources_from_json(source_json)
                if sources:
                    event = format_untrusted_tool_result(
                        "persisted routine sources: " + ", ".join(sources),
                        event,
                    )
            lines.append(f"- ID: {r['id']} | Event: {event} | Day: {r['day']} | Time: {r['time']}")
        return "\n".join(lines)
    except Exception as e:
        return f"Error searching routines: {str(e)}"

def _get_routine_names_for_intent_classification() -> list[str]:
    try:
        from memory.routine_db import get_connection
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT event_name FROM routines WHERE event_name IS NOT NULL AND event_name != ''")
        rows = cur.fetchall()
        conn.close()
        return [r[0] for r in rows if r and r[0]]
    except Exception:
        return []


def _looks_like_manual_followup_control(text: str) -> bool:
    normalized = unicodedata.normalize("NFKD", str(text or "").lower())
    normalized = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    control_markers = ("pending followup", "followup", "follow-up", *t("tools.system.control_markers_el"))
    action_markers = t("tools.system.action_markers_el")
    return any(m in normalized for m in control_markers) and any(m in normalized for m in action_markers)


@tool
def control_routine_notifications(event_name: str, action: str, until_date: str = "", source_text: str = "") -> str:
    """
    [OVERRIDE]: Manual control of a routine's proactive reminders, ONLY when
    {config.USER_NAME} EXPLICITLY requests it within the conversation (not automatically by you or by the
    scheduled job — this is a separate channel, the user takes control).

    VERY IMPORTANT — NEVER call this just because the user told you a piece of INFORMATION
    (e.g., "Kid1 is away at camp", "he returns in 9 days", "we went to the beach today").
    A piece of information IS NOT a request. DO NOT use this for daily context changes or daily presence updates (like "Partner is home today" or "we are together now"). The system handles daily presence automatically via dynamic context flags (Suppress). Call this ONLY when there is an explicit request to check notifications (or for explicit multi-day vacation/away blocks) — words/
    meaning like "don't send me", "mute", "leave alone", "stop notifications",
    "reactivate". If the user is simply informing you about something, reply normally in the
    conversation — DO NOT guess that they want to mute and DO NOT scan other routines "just in case".
    One call per routine explicitly requested — no repetition of the same call in the same turn.

    EXAMPLES OF EXPLICIT REQUESTS (only these patterns, do not generalize to every routine):
    - "No need to send me about the park until Kid1 returns on 26/6"
      → action="mute", until_date="2026-06-26"
    - "Leave the alarm alone all week, I'm on afternoon shift at work"
      → action="mute", until_date=<calculate it YOURSELF from the context, e.g., next Sunday>
    - "Reactivate the notifications for the park" or "Kid1 is back"
      → action="unmute"
    - "Don't send ANYTHING, not even a warm message, for the park until he returns"
      → action="silence_emotional"
    - "You can send some message for the park again while he is away"
      → action="allow_emotional"

    ARGUMENTS:
    - event_name: the name of the routine as spoken by the user (e.g., "park", "work
      alarm") — DOES NOT need to be exact, a fuzzy match is performed in memory.
    - action: one of "mute", "unmute", "silence_emotional", "allow_emotional".
    - until_date: ONLY for action="mute". Format YYYY-MM-DD. Calculate it YOURSELF from the
      context (today + X days, "this week" etc) — DO NOT ask the
      user to explicitly state it in ISO format.
    - source_text: The exact original message/sentence of the user that led to
      this call (ALWAYS MANDATORY).
    """
    from datetime import datetime
    from memory.routine_db import (
        find_routines_for_schedule_control, set_routine_muted_until, clear_routine_muted_until,
        set_sentimental_silenced, get_sentimental_info, get_routine_muted_until,
    )
    
    routine_names = _get_routine_names_for_intent_classification()
    intent_result = classify_routine_intent(source_text, routine_names=routine_names)
    if intent_result.intent == "context_update":
        return t("tools.system.context_update_not_notif")

    VALID_ACTIONS = {"mute", "unmute", "silence_emotional", "allow_emotional"}
    if action not in VALID_ACTIONS:
        return t("tools.system.invalid_action", action=action, allowed=", ".join(sorted(VALID_ACTIONS)))

    changed = 0
    already_ok = 0

    try:
        routines = find_routines_for_schedule_control(event_name)
    except Exception as e:
        return t("tools.system.routine_search_err", e=str(e))

    if not routines:
        return t("tools.system.routine_no_clear_match", event_name=event_name)

    results = []

    try:
        if action == "mute":
            until_date = (until_date or "").strip()
            if not until_date:
                return t("tools.system.req_until_date")
            try:
                datetime.strptime(until_date, "%Y-%m-%d")
            except ValueError:
                return t("tools.system.invalid_date_format", until_date=until_date)
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                existing_until = get_routine_muted_until(r_id)
                if existing_until and existing_until >= until_date:
                    results.append(t("tools.system.routine_already_muted", day=day, label=label, existing_until=existing_until))
                    already_ok += 1
                    continue
                set_routine_muted_until(r_id, until_date)
                results.append(t("tools.system.routine_muted", day=day, label=label, until_date=until_date))
                changed += 1
            
            if changed == 0 and already_ok > 0:
                return t("tools.system.routines_already_desired", event_name=event_name)
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)

        if action == "unmute":
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                clear_routine_muted_until(r_id)
                results.append(t("tools.system.routine_unmuted", day=day, label=label))
                changed += 1
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)

        if action == "silence_emotional":
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                info = get_sentimental_info(r_id)
                if not info["muted_until"]:
                    results.append(t("tools.system.routine_not_muted", day=day, label=label))
                    already_ok += 1
                    continue
                set_sentimental_silenced(r_id, True)
                results.append(t("tools.system.routine_silent_ok", day=day, label=label))
                changed += 1
            if changed == 0 and already_ok > 0:
                return t("tools.system.routines_already_desired", event_name=event_name)
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)

        if action == "allow_emotional":
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                set_sentimental_silenced(r_id, False)
                results.append(t("tools.system.routine_warm_ok", day=day, label=label))
                changed += 1
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)
    except Exception as e:
        return t("tools.system.routine_update_err2", e=str(e))

    return "❌ Unknown error."


@tool
def control_routine_condition(event_name: str, action: str, condition_type: str = "", payload_json: str = "", condition_mode: str = "", source_text: str = "", day_of_week: str = "", time_str: str = "") -> str:
    """
    [OVERRIDE]: Adds or Removes "conditions" from a routine.
    Use this WHEN the user requests a routine to depend on an external factor
    (e.g., shift, weather, location) or when they state that something "does not apply when..."

    ARGUMENTS:
    - event_name: The name of the target routine as spoken by the user — fuzzy matched.
    - action: "add" (to add a condition) or "remove" (to clear the condition).
    - condition_type: e.g., "shift_mode" (shift dependency), "context_flag" (dependency on a general flag like school_open).
    - payload_json: JSON string with the parameters (e.g., '{"flag": "user_out_of_home", "equals": false}').
      VALID context_flag KEYS: 'user_out_of_home', 'family_at_home', 'partner_with_user', 'kid1_away_from_home', 'user_at_work', 'kid1_with_user', 'kid1_with_partner', 'current_shift', 'football_season', 'school_open'.
      DO NOT INVENT NEW FLAG NAMES. For "at home", use {"flag": "user_out_of_home", "equals": false}.
    - condition_mode: "allow_when_true" (allowed ONLY if true) or "suppress_when_true" (CANCELLED when true).
    - source_text: The exact original message/sentence of the user (ALWAYS MANDATORY).
    - day_of_week: (Optional) If the user specified a day (e.g., "Sunday", "Monday").
    - time_str: (Optional) If the user specified a time (e.g., "13:00").


    EXAMPLES:
    "When I have an afternoon shift, my partner takes the park" (meaning the park for me does NOT apply in the afternoon)
    -> action="add", condition_type="shift_mode", payload_json='{"flag": "current_shift", "equals": "afternoon"}', condition_mode="suppress_when_true"

    "Training only applies when I have a morning shift"
    -> action="add", condition_type="shift_mode", payload_json='{"flag": "current_shift", "equals": "morning"}', condition_mode="allow_when_true"
    """
    import json
    from memory.routine_db import (
        find_routines_for_schedule_control, append_routine_condition, set_routine_condition
    )

    routine_names = _get_routine_names_for_intent_classification()
    intent_result = classify_routine_intent(source_text, routine_names=routine_names)
    is_structured_manual_condition = bool(
        action == "add" and condition_type and payload_json and condition_mode
    )
    if intent_result.intent == "context_update" and not is_structured_manual_condition:
        return t("tools.system.context_update_not_cond")

    VALID_ACTIONS = {"add", "remove"}
    if action not in VALID_ACTIONS:
        return t("tools.system.invalid_action_cond", action=action)

    routines = find_routines_for_schedule_control(event_name, day_of_week=day_of_week if day_of_week else None, time_str=time_str if time_str else None)
    if not routines:
        return f"❌ No routine found matching '{event_name}'."

    results = []
    changed = 0

    if action == "add":
        if not condition_type or not payload_json or not condition_mode:
            return t("tools.system.req_add_params")
        try:
            json.loads(payload_json) # Validate JSON
        except json.JSONDecodeError:
            return t("tools.system.invalid_payload")

        for routine in routines:
            r_id = routine["id"]
            label = routine["event"]
            r_day = str(routine["day"]).lower()
            
            # Smart check: If the condition concerns a shift and the routine is EXCLUSIVELY for the Weekend,
            # we ignore it automatically, as the user's shifts are Monday-Friday.
            # If the user explicitly provided day_of_week, then we allow it.
            if condition_type == "shift_mode" and r_day in ("saturday", "sunday") and not day_of_week:
                results.append(t("tools.system.routine_ignored_weekend", label=label, day=routine["day"]))
                continue

            added = append_routine_condition(r_id, condition_type=condition_type, condition_payload=payload_json, condition_mode=condition_mode, source_memory_ref="llm_agent")
            if added:
                results.append(t("tools.system.routine_cond_added", label=label, condition_type=condition_type, condition_mode=condition_mode))
                changed += 1
            else:
                results.append(t("tools.system.routine_cond_exists", label=label))

    elif action == "remove":
        for routine in routines:
            r_id = routine["id"]
            label = routine["event"]
            # Clear legacy conditions
            set_routine_condition(r_id, condition_type="", condition_payload="", condition_mode="", source_memory_ref="")
            # Clear conditions_json
            import sqlite3
            from memory.routine_db import get_connection, db_write_lock
            with db_write_lock:
                conn = get_connection(write=True)
                conn.execute("UPDATE routines SET conditions_json = NULL WHERE id = ?", (r_id,))
                conn.commit()
            results.append(t("tools.system.routine_cond_cleared", label=label))
            changed += 1

    if changed == 0:
        return t("tools.system.routines_no_change", event_name=event_name)
    return "\n".join(results)

@tool
def control_routine_schedule(event_name: str, action: str, until_date: str = "",
                              active_from: str = "", active_until: str = "",
                              resume_rule: str = "", reason: str = "", source_text: str = "") -> str:
    """
    [OVERRIDE]: Manual control of the SEASONAL/TEMPORARY inactivity of a routine
    (not notifications — that's what control_routine_notifications is for). Use
    it ONLY when {config.USER_NAME} EXPLICITLY asks to "freeze" / "stop" / "resume" a
    routine due to summer break, camp, season change, etc.

    DIFFERENCE FROM control_routine_notifications:
    - control_routine_notifications = "do not SEND me" (notification layer, the routine
      remains active in terms of confidence/missed-tracking).
    - control_routine_schedule = "this routine DOES NOT APPLY now" (business-logic layer —
      it does not enter missed/failed logic, confidence does not drop, it simply "freezes").
    For a summer break of a school/seasonal activity (e.g., football, school)
    ALWAYS use this tool, NOT mute, unless {config.USER_NAME} explicitly asks for "mute"/
    "mute notifications".

    VERY IMPORTANT — DO NOT call it just because the user gave you an INFORMATION (e.g.,
    "Kid1 is away at camp for 2 weeks"). An information IS NOT a request. Call
    it ONLY when there is an explicit request to pause/resume a routine.

    ARGUMENTS:
    - event_name: the name of the routine as spoken by the user — fuzzy match in memory.
    - action: one of "pause", "resume", "set_window", "clear_window".
      • pause → freezes the routine until until_date (YYYY-MM-DD), optional reason
        (e.g., "summer_break", "camp", "shift_change") and optional resume_rule
        (e.g., "every_september", "next_school_year", "manual_only").
      • resume → removes the pause, the routine immediately returns to normal operation.
      • set_window → sets active_from and/or active_until (YYYY-MM-DD) — the routine
        applies ONLY within this date window.
      • clear_window → removes the active_from/active_until window.
    - until_date: ONLY for action="pause". Calculate it YOURSELF from the context.
    - active_from / active_until: ONLY for action="set_window" (YYYY-MM-DD, only
      one of the two can be provided).
    - resume_rule: optional, along with action="pause" — how/when it will resume.
    - reason: optional, human description of the reason (e.g., "summer_break").
    - source_text: The exact original message/sentence of the user that led to
      this call (ALWAYS MANDATORY).

    EXAMPLE:
    "Kid1's football stops until September for the summer"
      → action="pause", until_date="2026-09-01", reason="summer_break",
        resume_rule="every_september"
    """
    from datetime import datetime
    from memory.routine_db import (
        find_routines_for_schedule_control, set_routine_paused_until, clear_routine_paused_until,
        set_routine_active_window, set_routine_resume_rule, get_routine_schedule_meta,
        normalize_event
    )
    routine_names = _get_routine_names_for_intent_classification()
    intent_result = classify_routine_intent(source_text, routine_names=routine_names)
    if intent_result.intent == "context_update":
        return t("tools.system.context_update_not_sched")

    VALID_ACTIONS = {"pause", "resume", "set_window", "clear_window"}
    if action not in VALID_ACTIONS:
        return t("tools.system.invalid_action", action=action, allowed=", ".join(sorted(VALID_ACTIONS)))

    def _valid_date(d: str) -> bool:
        try:
            datetime.strptime(d, "%Y-%m-%d")
            return True
        except ValueError:
            return False

    decision_text = source_text or event_name
    changed = 0
    already_ok = 0

    try:
        routines = find_routines_for_schedule_control(event_name)
    except Exception as e:
        return t("tools.system.routine_search_err", e=str(e))

    if not routines:
        return t("tools.system.routine_no_clear_match", event_name=event_name)

    results = []

    try:
        if action == "pause":
            until_date = (until_date or "").strip()
            if not until_date:
                return t("tools.system.req_until_date")
            if not _valid_date(until_date):
                return t("tools.system.invalid_date_format", until_date=until_date)
            reason_clean = reason.strip() or None
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                meta = get_routine_schedule_meta(r_id)
                existing_until = meta.get("paused_until")
                if existing_until and existing_until >= until_date:
                    results.append(t("tools.system.routine_already_frozen", day=day, label=label, existing_until=existing_until))
                    already_ok += 1
                    continue
                set_routine_paused_until(r_id, until_date, reason=reason_clean)
                if resume_rule.strip():
                    set_routine_resume_rule(r_id, resume_rule.strip())
                results.append(t("tools.system.routine_frozen", day=day, label=label, until_date=until_date, reason=t("tools.system.reason_prefix", reason=reason_clean) if reason_clean else ""))
                changed += 1
            if changed == 0 and already_ok > 0:
                return t("tools.system.routines_already_desired", event_name=event_name)
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)

        if action == "resume":
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                clear_routine_paused_until(r_id)
                results.append(t("tools.system.routine_unfrozen", day=day, label=label))
                changed += 1
            if changed == 0:
                return t("tools.system.routines_no_change", event_name=event_name)
            return "\n".join(results)

        if action == "set_window":
            active_from_clean = active_from.strip() or None
            active_until_clean = active_until.strip() or None
            if not active_from_clean and not active_until_clean:
                return t("tools.system.req_active_dates")
            if active_from_clean and not _valid_date(active_from_clean):
                return t("tools.system.invalid_active_from", active_from=active_from_clean)
            if active_until_clean and not _valid_date(active_until_clean):
                return t("tools.system.invalid_active_until", active_until=active_until_clean)
            reason_clean = reason.strip() or None
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                set_routine_active_window(r_id, active_from=active_from_clean, active_until=active_until_clean, reason=reason_clean)
                results.append(t("tools.system.routine_validity_window", day=day, label=label, active_from=active_from_clean or "—", active_until=active_until_clean or "—"))
            return "\n".join(results)

        if action == "clear_window":
            for routine in routines:
                r_id = routine["id"]
                label = routine["event"]
                day = routine.get("day") or "?"
                set_routine_active_window(r_id, active_from=None, active_until=None)
                results.append(t("tools.system.routine_validity_cleared", day=day, label=label))
            return "\n".join(results)
    except Exception as e:
        return t("tools.system.routine_sched_err", e=str(e))

    return "❌ Unknown error."


@tool
def control_routine_cooldown(
    event_name: str,
    action: str,
    source_text: str = "",
    day_of_week: str = "",
    time_str: str = "",
) -> str:
    """
    [OVERRIDE]: Manual routine cooldown control.

    Use this ONLY when the user explicitly requests something like:
    - "reset the cooldown"
    - "reset cooldown"
    - "take it out of cooldown"
    - "I want it to be sent normally tomorrow"
    - "to be resent normally"

    DO NOT call this for a simple context update.
    DO NOT call this just because the user was late to reply.
    DO NOT confuse this with mute/schedule/conditions.

    ARGUMENTS:
    - event_name: routine name as stated by the user
    - action: "reset" only
    - source_text: the exact message of the user
    - day_of_week/time_str: optional clarification if multiple similar routines exist
    """
    from memory.routine_db import (
        find_routines_for_schedule_control,
        reset_routine_cooldown,
        get_routine_notify_info,
    )
    routine_names = _get_routine_names_for_intent_classification()
    intent_result = classify_routine_intent(source_text, routine_names=routine_names)
    if intent_result.intent == "context_update":
        return "ℹ️ This looks like a context/fact update, not a manual routine cooldown override command. Cooldown not reset."

    VALID_ACTIONS = {"reset"}
    if action not in VALID_ACTIONS:
        return f"❌ Invalid action: '{action}'. Allowed: reset."

    routines = find_routines_for_schedule_control(
        event_name,
        day_of_week=day_of_week if day_of_week else None,
        time_str=time_str if time_str else None,
    )
    if not routines:
        return f"❌ No routine found matching '{event_name}'."

    results = []
    changed = 0

    for routine in routines:
        r_id = routine["id"]
        label = routine["event"]
        r_day = routine.get("day") or "?"
        r_time = routine.get("time") or "?"
        before = get_routine_notify_info(r_id)

        reset_routine_cooldown(r_id, clear_last_notified=True)

        after = get_routine_notify_info(r_id)
        results.append(
            f"🔄 [{r_day} {r_time}] Routine '{label}' removed from cooldown "
            f"({before['cooldown_hours']}h -> {after['cooldown_hours']}h)."
        )
        changed += 1

    if changed == 0:
        return f"ℹ️ No cooldown change for: {event_name}"

    return "\n".join(results)


@tool
def control_pending_followup(
    subject_query: str,
    action: str,
    source_text: str = "",
    topic: str = "",
    delay_minutes: int = 0,
    target_window: str = "",
) -> str:
    """
    [OVERRIDE]: Manual check of pending conversational follow-ups.

    Use this ONLY when the user explicitly asks to change/delete/defer a pending follow-up.

    Examples:
    - "delete the pending followup for the steaks"
    - "change the pending for the park to later"
    - "defer the followup for Sophia to tomorrow afternoon"
    - "fix the old pending followups"

    actions:
    - "delete": deletes the matching follow-up
    - "defer": defers the matching follow-up with new minutes/window
    - "repair_legacy": backfills old followups that are missing metadata fields
    """
    from memory.pending_followups import (
        backfill_legacy_followups,
        delete_followup,
        defer_followup,
        find_followups_for_control,
        find_pending_followups,
    )

    action = (action or "").strip().lower()
    if action not in {"delete", "defer", "repair_legacy"}:
        return "❌ Invalid action. Allowed: delete, defer, repair_legacy."

    if action != "repair_legacy":
        if not _looks_like_manual_followup_control(source_text):
            return (
                "ℹ️ This looks more like a context update rather than a manual "
                "pending follow-up override command. No changes made."
            )

    if action == "repair_legacy":
        repaired = backfill_legacy_followups(force_retime=True)
        rows = find_pending_followups(limit=10)
        if not repaired:
            return "ℹ️ No legacy pending followups needed repair."
        preview = []
        for row in rows[:5]:
            preview.append(
                f"- #{row['id']} {row['subject']} -> {row['followup_after_ts']}"
            )
        body = "\n".join(preview)
        return f"🛠️ Repaired {repaired} legacy pending followups.\n{body}"

    matches = find_followups_for_control(subject_query, topic=topic)
    if not matches:
        return f"ℹ️ No pending/sent follow-up found matching '{subject_query}'."

    if len(matches) > 1:
        opts = "\n".join(
            f"- #{m['id']} {m['subject']} ({m['topic']}, {m['status']})"
            for m in matches[:5]
        )
        return f"⚠️ Found multiple pending followups. Please specify:\n{opts}"

    item = matches[0]
    if action == "delete":
        ok = delete_followup(item["id"], reason="manual_delete")
        if not ok:
            return f"❌ Failed to delete pending follow-up #{item['id']}."
        return f"✅ Deleted pending follow-up #{item['id']} for '{item['subject']}'."

    if delay_minutes <= 0 and not target_window.strip():
        return "❌ Defer requires delay_minutes and/or target_window."

    defer_followup(
        item["id"],
        delay_minutes=delay_minutes or int(item.get("metadata", {}).get("delay_minutes_final") or 60),
        reason="manual_defer",
        target_window=(target_window or str(item.get("metadata", {}).get("target_window") or "")).strip(),
        topic=item["topic"],
    )
    updated = find_pending_followups(limit=20)
    refreshed = next((row for row in updated if row["id"] == item["id"]), None)
    if not refreshed:
        return f"✅ Pending follow-up #{item['id']} deferred."
    return (
        f"✅ Pending follow-up #{item['id']} for '{refreshed['subject']}' deferred.\n"
        f"New due: {refreshed['followup_after_ts']}\n"
        f"New expiry: {refreshed['expires_at']}"
    )


@tool
def manage_list(
    action: str,
    list_name: str,
    item: str = "",
    external_content_sources_json: str = "",
) -> str:
    """Manages lists. Actions: 'add', 'remove', 'read', 'clear', 'delete'.
    For multiple items at once, separate them with a comma (item='milk, cheese').
    For destructive actions ('clear', 'delete'), item must be '__CONFIRMED_CLEAR__'.
    external_content_sources_json is internal approval provenance. Do not set it manually."""
    if action in {"clear", "delete"} and item != "__CONFIRMED_CLEAR__":
        return (
            f"Error: Refusing to {action} list '{list_name}' without explicit confirmation token."
        )
    conn = None
    try:
        conn = sqlite3.connect(STATE_DB)
        cursor = conn.cursor()
        
        cursor.execute("SELECT DISTINCT list_name FROM lists")
        existing_lists = [row[0] for row in cursor.fetchall()]
        
        if list_name not in existing_lists:
            list_name_lower = list_name.lower()
            for existing_key in existing_lists:
                if list_name_lower in existing_key.lower() or existing_key.lower().startswith(list_name_lower):
                    list_name = existing_key
                    break

        if action == "read":
            cursor.execute(
                "SELECT item, external_content_sources_json FROM lists WHERE list_name=?",
                (list_name,),
            )
            rows = cursor.fetchall()
            items = [row[0] for row in rows]
            if not items:
                return f"The list '{list_name}' is empty."
            from core.untrusted_content import (
                external_content_sources_from_json,
                format_untrusted_tool_result,
            )

            rendered_items = []
            for list_item, raw_sources in rows:
                sources = external_content_sources_from_json(raw_sources or "")
                if sources:
                    list_item = format_untrusted_tool_result(
                        "persisted list sources: " + ", ".join(sources),
                        list_item,
                    )
                rendered_items.append(f"- {list_item}")
            return f"Contents of '{list_name}':\n" + "\n".join(rendered_items)

        to_process = [i.strip() for i in item.split(",")] if item else []
        from core.untrusted_content import external_content_sources_from_json

        external_sources = external_content_sources_from_json(external_content_sources_json)
        provenance_json = json.dumps(external_sources)

        if action == "add":
            for obj in to_process:
                if obj:
                    cursor.execute(
                        "SELECT id, external_content_sources_json FROM lists "
                        "WHERE list_name=? AND item=?",
                        (list_name, obj),
                    )
                    existing = cursor.fetchone()
                    if existing is None:
                        cursor.execute(
                            "INSERT INTO lists (list_name, item, external_content_sources_json) "
                            "VALUES (?, ?, ?)",
                            (list_name, obj, provenance_json),
                        )
                    elif external_sources:
                        existing_sources = external_content_sources_from_json(existing[1] or "")
                        merged_sources = sorted(set(existing_sources) | set(external_sources))
                        cursor.execute(
                            "UPDATE lists SET external_content_sources_json=? WHERE id=?",
                            (json.dumps(merged_sources), existing[0]),
                        )
        elif action == "remove":
            for obj in to_process:
                cursor.execute("DELETE FROM lists WHERE list_name=? AND item=?", (list_name, obj))
        elif action == "clear" or action == "delete":
            cursor.execute("DELETE FROM lists WHERE list_name=?", (list_name,))

        conn.commit()

        added_str = ", ".join(to_process) if to_process else "none"
        return f"System: Action '{action}' completed (Items: {added_str})."
    except Exception as e:
        return f"Error: List error: {str(e)}"
    finally:
        if conn:
            conn.close()


# ────────────────────────────────────────────────────────────────
# GOOGLE SERVICES
# ────────────────────────────────────────────────────────────────

from astakos_skills.gcalendar import google_calendar_tool


def _normalize_google_task_due(due: str | None) -> str | None:
    if not due:
        return None
    due = due.strip()
    if len(due) == 10:
        return f"{due}T00:00:00.000Z"
    return due


@tool
def google_tasks_tool(
    action: str,
    title: str = "",
    due: str = None,
    task_id: str = None,
    notes: str = "",
    tasklist_id: str = "@default",
    max_results: int = 20,
) -> str:
    """
    Manages Google Tasks.
    Actions:
      'create'   — new task (requires title, optionally due/notes)
      'list'     — list of open tasks
      'complete' — complete a task (requires task_id)
      'update'   — modify title/due/notes (requires task_id and at least one field)
      'delete'   — delete a task (requires task_id, CRITICAL approval)
    🚨 [MASTER-RULE FOR TITLE]: The 'title' parameter MUST clearly describe the task (e.g., 'Rose bush care').
    It is STRICTLY FORBIDDEN to use command verbs (e.g., 'put', 'do', 'remind me', 'reminder') as a title.
    If the user simply says 'add a reminder' without specifying the topic, DO NOT call this tool. Ask them first what they want you to write!
    """
    try:
        action = (action or "list").strip().lower()
        tasklist_id = tasklist_id or "@default"
        print(f"\033[93m[Tasks]: Action '{action}'...\033[0m")
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, ['https://www.googleapis.com/auth/tasks'])
        service = build('tasks', 'v1', credentials=creds)

        if action == "list":
            result = service.tasks().list(
                tasklist=tasklist_id,
                showCompleted=False,
                maxResults=max(1, min(int(max_results or 20), 50)),
            ).execute()
            items = result.get("items", [])
            if not items:
                return "✅ No open Google Tasks."
            lines = ["📋 Open Google Tasks:"]
            for task in items:
                due_text = task.get("due", "")[:10]
                due_part = f" | due: {due_text}" if due_text else ""
                lines.append(f"• {task.get('title', t('tools.system.task_no_title'))} | ID: `{task.get('id')}`{due_part}")
            return "\n".join(lines)

        if action == "create":
            if not title:
                return "❌ Create requires a title."
            task = {"title": title}
            normalized_due = _normalize_google_task_due(due)
            if normalized_due:
                task["due"] = normalized_due
            if notes:
                task["notes"] = notes
            created = service.tasks().insert(tasklist=tasklist_id, body=task).execute()
            return f"✅ Task '{created.get('title', title)}' added to Google Tasks! ID: `{created.get('id')}`"

        if action == "complete":
            if not task_id:
                return "❌ Complete requires a task_id."
            service.tasks().patch(
                tasklist=tasklist_id,
                task=task_id,
                body={"status": "completed"},
            ).execute()
            return f"✅ Google Task `{task_id}` completed."

        if action == "update":
            if not task_id:
                return "❌ Update requires a task_id."
            body = {}
            if title:
                body["title"] = title
            normalized_due = _normalize_google_task_due(due)
            if normalized_due:
                body["due"] = normalized_due
            if notes:
                body["notes"] = notes
            if not body:
                return "❌ Update requires title, due, or notes."
            updated = service.tasks().patch(tasklist=tasklist_id, task=task_id, body=body).execute()
            return f"✅ Google Task updated: {updated.get('title', task_id)}"

        if action == "delete":
            if not task_id:
                return "❌ Delete requires a task_id."
            service.tasks().delete(tasklist=tasklist_id, task=task_id).execute()
            return f"🗑️ Google Task `{task_id}` deleted."

        return "❌ Unknown action. Try: list, create, complete, update, delete."
    except Exception as e:
        return f"Tasks Error: {str(e)}"

@tool
def create_file_tool(file_type: str, filename: str, data: str) -> str:
    """
    [WARNING: For .docx, .xlsx, .pptx files, the use of this tool is FORBIDDEN. Use ONLY run_officecli]
    Creates local PDF, TXT (and legacy DOCX/XLSX) files.
    file_type: 'docx', 'pdf', 'xlsx', 'txt'
    filename: The name of the file (e.g., 'report.txt')
    data: The content. 
    """
    import os
    import json
    from config import BASE_DIR

    output_dir = os.path.realpath(os.path.join(BASE_DIR, "outputs"))
    os.makedirs(output_dir, exist_ok=True)

    # [SECURITY]: basename + resolve check — prevents path traversal (e.g., ../config.py)
    safe_filename = os.path.basename(filename)
    if not safe_filename:
        return "❌ Error: Invalid filename."
    full_path = os.path.realpath(os.path.join(output_dir, safe_filename))
    if not full_path.startswith(output_dir + os.sep) and full_path != output_dir:
        return "❌ Error: Path outside outputs is not allowed."
    file_type = file_type.lower()

    try:
        if file_type == "docx":
            import docx
            doc = docx.Document()
            for line in data.split("\n"):
                doc.add_paragraph(line)
            doc.save(full_path)

        elif file_type == "xlsx":
            import pandas as pd
            try:
                content = json.loads(data)
                df = pd.DataFrame(content)
            except (json.JSONDecodeError, ValueError):
                df = pd.DataFrame([data], columns=["Content"])
            df.to_excel(full_path, index=False)

        elif file_type == "pdf":
            from fpdf import FPDF
            font_path = os.path.join(BASE_DIR, "assets", "DejaVuSans.ttf")
            pdf = FPDF()
            pdf.add_page()
            if os.path.exists(font_path):
                pdf.add_font("DejaVu", "", font_path, uni=True)
                pdf.set_font("DejaVu", size=12)
            else:
                # Fallback without Greek if the font is missing
                pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, data)
            pdf.output(full_path)

        elif file_type in ["txt", "json", "csv", "html", "md"]:
            with open(full_path, "w", encoding="utf-8") as f:
                f.write(data)

        else:
            return f"❌ Error: Type '{file_type}' not supported."

        return f"✅ Ready Boss! File created successfully.\n[CREATED_FILE: {full_path}]"

    except Exception as e:
        return f"❌ Error during creation: {str(e)}"
@tool
def generate_image_tool(prompt: str) -> str:
    """
    Creates an image based on a description (prompt) via Vertex AI Imagen.
    """
    import os
    import time
    from slugify import slugify
    from config import BASE_DIR

    output_dir = os.path.join(BASE_DIR, "outputs")
    os.makedirs(output_dir, exist_ok=True)

    safe_filename = slugify(prompt[:30]) or "gen_image"
    filename = f"{safe_filename}_{int(time.time())}.jpg"
    full_path = os.path.join(output_dir, filename)

    # ── Vertex AI Imagen ──────────────────────────────────────────
    try:
        from google import genai
        from google.genai import types

        api_key = config.GEMINI_API_KEY
        if api_key:
            client = genai.Client(api_key=api_key)
        else:
            client = genai.Client(
                vertexai=True,
                project=config.PROJECT_ID,
                location=config.LOCATION or "us-central1"
            )
            
        response = client.models.generate_images(
            model='imagen-3.0-generate-001',
            prompt=prompt,
            config=types.GenerateImagesConfig(
                number_of_images=1,
                aspect_ratio="1:1"
            )
        )
        
        if not response.generated_images:
            return "❌ Vertex AI Imagen returned no image."
            
        response.generated_images[0].image.save(full_path)
        return f"✅ Ready! Image created.\n[SEND_PHOTO: {full_path}]"

    except Exception as e:
        return f"❌ Error Vertex AI Imagen: {str(e)}"
def _escape_drive_query_value(value: str) -> str:
    return value.replace("\\", "\\\\").replace("'", "\\'")


@tool
def drive_manager(
    action: str = "list_files",
    file_id: str = None,
    local_path: str = None,
    folder_id: str = "12YrIZ3uAQWmmwIlEkIkDf-4gcz2P8Ktv",
    query: str = None,
    new_name: str = None,
    target_folder_id: str = None,
    share_email: str = None,
    share_role: str = "reader",
) -> str:
    """Manages {config.USER_NAME}'s Google Drive.

    Actions:
      'list_files'   — List files in folder (default: root astakos folder)
      'search'       — Search by name or keyword (requires query=)
      'download'     — Download file (requires file_id=)
      'upload'       — Upload file (requires local_path=)
      'delete'       — Delete file (requires file_id=)
      'rename'       — Rename (requires file_id= + new_name=)
      'move'         — Move to another folder (requires file_id= + target_folder_id=)
      'share'        — Share (requires file_id= + share_email= + share_role='reader'/'writer')
      'create_folder'— Create folder (requires new_name=, optionally folder_id= for parent)
      'info'         — File information (requires file_id=)
    """
    try:
        from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
        import io

        action = (action or "list_files").strip().lower()
        print(f"\033[93m[Drive]: Action '{action}'...\033[0m")
        creds   = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)
        service = build('drive', 'v3', credentials=creds)

        # ── LIST FILES ───────────────────────────────────────────
        if action == "list_files":
            results = service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                fields="files(id, name, mimeType, size, modifiedTime)",
                orderBy="modifiedTime desc",
                pageSize=50
            ).execute()
            items = results.get('files', [])
            if not items:
                return "📁 The folder is empty."
            lines = ["📁 Files in Drive:\n"]
            for i in items:
                size_kb = round(int(i.get('size', 0)) / 1024, 1) if i.get('size') else "—"
                mod = i.get('modifiedTime', '')[:10]
                lines.append(f"• {i['name']} | ID: `{i['id']}` | {size_kb} KB | {mod}")
            return "\n".join(lines)

        # ── SEARCH ───────────────────────────────────────────────
        elif action == "search":
            if not query:
                return "❌ Requires query= for search."
            q_str = f"name contains '{_escape_drive_query_value(query)}' and trashed=false"
            results = service.files().list(
                q=q_str,
                fields="files(id, name, mimeType, size, modifiedTime, parents)",
                orderBy="modifiedTime desc",
                pageSize=20
            ).execute()
            items = results.get('files', [])
            if not items:
                return f"🔍 No files found for '{query}'."
            lines = [f"🔍 Results for '{query}':\n"]
            for i in items:
                size_kb = round(int(i.get('size', 0)) / 1024, 1) if i.get('size') else "—"
                mod = i.get('modifiedTime', '')[:10]
                lines.append(f"• {i['name']} | ID: `{i['id']}` | {size_kb} KB | {mod}")
            return "\n".join(lines)

        # ── DOWNLOAD ─────────────────────────────────────────────
        elif action == "download":
            if not file_id:
                return "❌ Requires file_id=."
            file_metadata = service.files().get(fileId=file_id, fields="name,mimeType").execute()
            mime_type = file_metadata.get('mimeType', '')
            file_name = file_metadata.get('name', 'downloaded_file')

            # Google Docs/Sheets/Slides → export as text/xlsx/pptx
            export_map = {
                'application/vnd.google-apps.document':     ('text/plain', '.txt'),
                'application/vnd.google-apps.spreadsheet':  ('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', '.xlsx'),
                'application/vnd.google-apps.presentation': ('application/vnd.openxmlformats-officedocument.presentationml.presentation', '.pptx'),
            }
            if mime_type in export_map:
                export_mime, ext = export_map[mime_type]
                request = service.files().export_media(fileId=file_id, mimeType=export_mime)
                if not file_name.endswith(ext):
                    file_name += ext
            else:
                request = service.files().get_media(fileId=file_id)

            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()

            # [SECURITY]: always download to outputs — local_path is ignored if it is outside
            from config import BASE_DIR as _BASE_DIR
            _outputs_dir = os.path.realpath(os.path.join(_BASE_DIR, "outputs"))
            os.makedirs(_outputs_dir, exist_ok=True)
            if local_path:
                _lp_real = os.path.realpath(local_path)
                if not _lp_real.startswith(_outputs_dir + os.sep):
                    return f"❌ Forbidden download path: only within outputs/ allowed."
                save_target = _lp_real
            else:
                save_target = os.path.join(_outputs_dir, os.path.basename(file_name))
            os.makedirs(os.path.dirname(save_target), exist_ok=True)
            with open(save_target, "wb") as f:
                f.write(fh.getvalue())

            # If it is text, also return the content
            if mime_type == 'application/vnd.google-apps.document' or file_name.endswith('.txt'):
                content = fh.getvalue().decode('utf-8', errors='ignore')[:6000]
                return f"✅ '{file_name}' downloaded → {save_target}\n\n{content}"
            return f"✅ '{file_name}' downloaded → {save_target}"

        # ── UPLOAD ───────────────────────────────────────────────
        elif action == "upload":
            if not local_path or not os.path.exists(local_path):
                return f"❌ File not found: {local_path}"
            # [SECURITY]: upload only from allowed dirs — prevents uploading credentials/config
            from config import BASE_DIR as _BASE_DIR
            _upload_allowed = [
                os.path.realpath(os.path.join(_BASE_DIR, "outputs")),
                os.path.realpath(os.path.join(_BASE_DIR, "telegram_uploads")),
                os.path.realpath(os.path.join(_BASE_DIR, "telegram_photos")),
                os.path.realpath(os.path.join(_BASE_DIR, "watch_folder")),
            ]
            _lp_real = os.path.realpath(local_path)
            if not any(_lp_real.startswith(d + os.sep) or _lp_real == d for d in _upload_allowed):
                return f"❌ Forbidden upload path: only from outputs/, telegram_uploads/, telegram_photos/, watch_folder/ allowed."
            file_metadata = {'name': os.path.basename(local_path), 'parents': [folder_id]}
            media = MediaFileUpload(local_path, resumable=True)
            file = service.files().create(body=file_metadata, media_body=media, fields='id,name').execute()
            return f"✅ '{file.get('name')}' uploaded! (ID: {file.get('id')})"

        # ── DELETE ───────────────────────────────────────────────
        elif action == "delete":
            if not file_id:
                return "❌ Requires file_id=."
            meta = service.files().get(fileId=file_id, fields="name").execute()
            service.files().update(fileId=file_id, body={"trashed": True}).execute()
            return f"🗑️ '{meta.get('name')}' moved to trash."

        # ── RENAME ───────────────────────────────────────────────
        elif action == "rename":
            if not file_id or not new_name:
                return "❌ Requires file_id= and new_name=."
            service.files().update(fileId=file_id, body={"name": new_name}).execute()
            return f"✏️ Renamed to '{new_name}'."

        # ── MOVE ─────────────────────────────────────────────────
        elif action == "move":
            if not file_id or not target_folder_id:
                return "❌ Requires file_id= and target_folder_id=."
            file = service.files().get(fileId=file_id, fields="parents").execute()
            old_parents = ",".join(file.get('parents', []))
            service.files().update(
                fileId=file_id,
                addParents=target_folder_id,
                removeParents=old_parents,
                fields="id, parents"
            ).execute()
            return f"📦 File moved to folder {target_folder_id}."

        # ── SHARE ────────────────────────────────────────────────
        elif action == "share":
            if not file_id or not share_email:
                return "❌ Requires file_id= and share_email=."
            if share_role not in {"reader", "writer", "commenter"}:
                return "❌ share_role must be reader, writer, or commenter."
            permission = {"type": "user", "role": share_role, "emailAddress": share_email}
            service.permissions().create(fileId=file_id, body=permission, sendNotificationEmail=False).execute()
            return f"🔗 Shared with {share_email} as {share_role}."

        # ── CREATE FOLDER ─────────────────────────────────────────
        elif action == "create_folder":
            if not new_name:
                return "❌ Requires new_name= for folder name."
            metadata = {
                "name": new_name,
                "mimeType": "application/vnd.google-apps.folder",
                "parents": [folder_id]
            }
            folder = service.files().create(body=metadata, fields="id, name").execute()
            return f"📁 Folder '{folder.get('name')}' created (ID: {folder.get('id')})."

        # ── INFO ─────────────────────────────────────────────────
        elif action == "info":
            if not file_id:
                return "❌ Requires file_id=."
            meta = service.files().get(
                fileId=file_id,
                fields="name,mimeType,size,modifiedTime,createdTime,parents,webViewLink,owners"
            ).execute()
            size_kb = round(int(meta.get('size', 0)) / 1024, 1) if meta.get('size') else "—"
            owners = ", ".join(o.get('emailAddress','') for o in meta.get('owners', []))
            return (
                f"📄 *{meta.get('name')}*\n"
                f"Type: {meta.get('mimeType')}\n"
                f"Size: {size_kb} KB\n"
                f"Created: {meta.get('createdTime','')[:10]}\n"
                f"Modified: {meta.get('modifiedTime','')[:10]}\n"
                f"Owner: {owners}\n"
                f"Link: {meta.get('webViewLink','—')}"
            )

        return "❌ Unknown action. See docstring for options."

    except Exception as e:
        return f"❌ Drive Error: {str(e)}"


# ────────────────────────────────────────────────────────────────
# FILE & DEV TOOLS
# ────────────────────────────────────────────────────────────────

@tool
def read_local_file(file_path: str) -> str:
    """Reads PDF, XLSX, CSV, DOCX, TXT, PY, JS (Mastro-Optimized)."""
    import os
    from config import BASE_DIR, PHOTOS_DIR
    
    # Path cleanup
    file_path = file_path.strip().strip("'").strip('"')
    filename = os.path.basename(file_path)
    base_dir = BASE_DIR

    # [SECURITY]: Only these folders are allowed to be read
    _allowed_dirs = [
        os.path.realpath(PHOTOS_DIR),
        os.path.realpath(os.path.join(base_dir, "telegram_uploads")),
        os.path.realpath(os.path.join(base_dir, "telegram_photos")),
        os.path.realpath(os.path.join(base_dir, "uploads")),
        os.path.realpath(os.path.join(base_dir, "outputs")),
        os.path.realpath(os.path.join(base_dir, "watch_folder")),
        # [SELF-DIAGNOSIS]: Source code folders — Lobster can read
        # its code for self-debugging (e.g., why a tool failed).
        os.path.realpath(os.path.join(base_dir, "tools")),
        os.path.realpath(os.path.join(base_dir, "core")),
        os.path.realpath(os.path.join(base_dir, "memory")),
        os.path.realpath(os.path.join(base_dir, "services")),
        os.path.realpath(os.path.join(base_dir, "clients")),
        os.path.realpath(os.path.join(base_dir, "astakos_skills")),
        os.path.realpath(os.path.join(base_dir, "api")),
    ]
    _allowed_files = [
        os.path.realpath(os.path.join(BASE_DIR, "messenger_draft.json")),
        os.path.realpath(os.path.join(BASE_DIR, "linkedin_draft.json")),
    ]
    # [SECURITY]: Sensitive files that are not allowed even if they are in an allowed dir
    _blocked_filenames = {
        "config.py", ".env", "secrets.py",
    }
    _blocked_extensions = {".db", ".sqlite", ".sqlite3", ".key", ".pem"}

    def _is_blocked(path):
        name = os.path.basename(path)
        ext  = os.path.splitext(name)[1].lower()
        return name in _blocked_filenames or ext in _blocked_extensions

    def _in_allowed(path):
        real = os.path.realpath(path)
        if _is_blocked(path):
            return False
        return any(real.startswith(d + os.sep) or real == d for d in _allowed_dirs)

    def _is_allowed_file(path):
        real = os.path.realpath(path)
        if _is_blocked(path):
            return False
        return any(real == f for f in _allowed_files)

    full_path = None
    print(f"\033[93m[Tool Debug]: Searching for file: {filename}\033[0m")

    # If an absolute path was provided, check that it is within the allowed dirs
    if os.path.isabs(file_path):
        if os.path.exists(file_path) and os.path.isfile(file_path) and (_in_allowed(file_path) or _is_allowed_file(file_path)):
            full_path = file_path
            print(f"\033[92m[Tool Debug]: ✅ Absolute path within allowed -> {full_path}\033[0m")
        elif os.path.exists(file_path):
            return f"❌ Forbidden path: {os.path.basename(file_path)} is outside approved folders."

    # Exact allowlist for root-level runtime files such as the Messenger draft.
    if not full_path:
        for allowed_file in _allowed_files:
            if filename == os.path.basename(allowed_file) and os.path.exists(allowed_file) and os.path.isfile(allowed_file):
                full_path = allowed_file
                print(f"\033[92m[Tool Debug]: ✅ Exact allowed file -> {full_path}\033[0m")
                break

    # Search by basename in the allowed dirs_
    if not full_path:
        for d in _allowed_dirs:
            test_path = os.path.join(d, filename)
            if os.path.exists(test_path) and os.path.isfile(test_path) and _in_allowed(test_path):
                full_path = test_path
                print(f"\033[92m[Tool Debug]: ✅ Found at -> {full_path}\033[0m")
                break

    if not full_path:
        return f"❌ Error: File {filename} not found in search folders."

    ext = os.path.splitext(full_path)[1].lower()

    try:
        if ext == ".pdf":
            text = ""
            reader = PdfReader(full_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
                if len(text) > 12000: # Limit to avoid "choking" the context
                    break
            
            if not text.strip():
                return f"⚠️ The PDF ({filename}) seems to be scanned (image). Requires OCR to read."
                
            return f"📄 PDF ({filename}):\n{text[:12000]}"

        elif ext in [".xlsx", ".xls"]:
            import pandas as pd
            excel_file = pd.ExcelFile(full_path)
            output_text = f"📊 Excel ({filename}) - Sheets: {', '.join(excel_file.sheet_names)}\n\n"
            for sheet in excel_file.sheet_names:
                df = pd.read_excel(full_path, sheet_name=sheet).fillna("-")
                output_text += f"═══ Sheet: {sheet} ═══\n"
                output_text += df.head(50).to_string(index=False) + "\n\n"
                if len(output_text) > 12000: break
            return output_text[:12000]

        elif ext == ".csv":
            import pandas as pd
            df = pd.read_csv(full_path).fillna("-")
            return f"📊 CSV ({filename}):\n{df.head(100).to_string(index=False)}"

        elif ext == ".docx":
            import docx
            doc = docx.Document(full_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return f"📝 Word ({filename}):\n{text[:12000]}"

        else: # TXT, PY, JS, etc.
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                return f"📄 File ({filename}):\n{f.read(12000)}"

    except Exception as e:
        return f"❌ Read error {filename}: {str(e)}"

@tool
def write_code(filename: str, code: str) -> str:
    """Writes code ONLY inside the astakos_skills folder."""
    safe_filename = os.path.basename(filename)
    if safe_filename in PROTECTED_FILES:
        return f"System Error: FORBIDDEN to modify {safe_filename}."

    if re.search(r"(^|\n)\s*@tool\b|langchain_core\.tools\s+import\s+tool", code):
        return (
            "System Error: skill tools must be created with write_custom_tool, "
            "then registered with register_tool dry_run/apply."
        )

    for word in DANGEROUS_WORDS:
        if word in code:
            return f"System Error: Code rejected ({word})."

    file_path = os.path.join(WORKSPACE_DIR, safe_filename)

    try:
        print(f"\033[93m[Dev]: Saving in {file_path}...\033[0m")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(code)
        return f"System: Code written to {file_path}."
    except Exception as e:
        return f"Write Error: {str(e)}"


@tool
def run_code(filename: str, script_args: str = "") -> str:
    """
    Executes a Python file ONLY from the astakos_skills folder.
    You can optionally pass arguments (script_args) as a string.
    Example: script_args="SKG KUT 2026-08-09 -r 2026-08-15"
    """
    import os
    import sys
    import subprocess
    from core.safe_executor import safe_execute

    safe_filename = os.path.basename(filename)
    file_path = os.path.join(WORKSPACE_DIR, safe_filename)

    if not os.path.exists(file_path):
        return f"Error: File {file_path} does not exist in Sandbox."

    # ── SafeExec check ───────────────────────────────────────────
    cmd_str = f"python {safe_filename} {script_args}".strip()
    check = safe_execute(cmd_str, lambda c: {"status": "ok"})
    if check.get("status") == "blocked":
        return f"🛡️ [SAFE EXECUTOR - BLOCKED]: {check['reason']}"
    if check.get("status") == "cancelled":
        return f"⚠️ [SAFE EXECUTOR]: Execution requires confirmation. Send again with `/confirm {cmd_str}`"
    # ────────────────────────────────────────────────────────────

    try:
        cmd = [sys.executable, file_path]
        if script_args:
            cmd.extend(script_args.split())

        print(f"\033[93m[Dev]: Executing {safe_filename} with arguments: {script_args}\033[0m")

        res = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
        output = res.stdout if res.stdout else ""
        if res.stderr:
            output += f"\nERRORS:\n{res.stderr}"

        return f"Terminal Output:\n{output[:5000]}" if output else "Executed successfully (no output)."

    except subprocess.TimeoutExpired:
        return "Error: Script hung (>20 seconds) and was terminated."
    except Exception as e:
        return f"Run Error: {str(e)}"


_SAFE_TOOL_MODULES = frozenset({
    "math", "json", "datetime", "re", "typing", "pydantic",
    "random", "string", "collections", "itertools", "decimal",
    "hashlib", "uuid", "time", "functools", "inspect",
    "langchain_core", "langchain_core.tools", "base64", "zlib",
    "gzip", "enum", "dataclasses", "copy",
})

_ALLOWED_EXTERNAL_TOOL_MODULES = frozenset({
    "httpx", "requests", "urllib", "bs4", "beautifulsoup4",
    "csv", "pandas", "sqlite3",
})

_FORBIDDEN_TOOL_MODULES = frozenset({
    "os", "sys", "subprocess", "ctypes", "shutil", "importlib",
    "builtins", "socket", "ftplib", "smtplib", "paramiko",
    "pickle", "shelve", "marshal", "pty", "commands", "posix",
    "nt", "signal", "threading", "multiprocessing", "asyncio", "gc",
})

_FORBIDDEN_TOOL_CALLS = frozenset({
    "eval", "exec", "open", "compile", "__import__",
    "globals", "locals", "vars", "getattr", "setattr", "delattr",
    "exit", "quit",
})

_FORBIDDEN_TOOL_DUNDERS = frozenset({
    "__builtins__", "__class__", "__bases__", "__subclasses__",
    "__globals__", "__code__", "__dict__", "__getattribute__",
    "__reduce__", "__reduce_ex__",
})


def _validate_custom_tool_ast(code: str, tool_name: str) -> tuple[bool, str, set[str]]:
    """Validate dynamically generated tool Python code via AST inspection."""
    if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", tool_name):
        return False, "System Error: invalid tool_name. Use a Python identifier, e.g. my_tool.", set()

    try:
        tree = ast.parse(code)
    except SyntaxError as se:
        return False, f"❌ Syntax error (line {se.lineno}): {se.msg}\nLook: {se.text}", set()
    except Exception as exc:
        return False, f"❌ Code parsing error: {exc}", set()

    detected_capabilities = set()

    for node in ast.walk(tree):
        # 1. Inspect imports
        if isinstance(node, ast.Import):
            for alias in node.names:
                root_module = alias.name.split(".")[0]
                if root_module in _FORBIDDEN_TOOL_MODULES:
                    return False, f"System Error: Rejected — forbidden module import: `{alias.name}`.", set()
                if root_module in _ALLOWED_EXTERNAL_TOOL_MODULES:
                    detected_capabilities.add(root_module)
                elif root_module not in _SAFE_TOOL_MODULES:
                    return False, f"System Error: Rejected — unapproved module import: `{alias.name}`.", set()

        elif isinstance(node, ast.ImportFrom):
            if node.level and node.level > 0:
                return False, "System Error: Rejected — relative imports are not permitted in skills.", set()
            if node.module:
                root_module = node.module.split(".")[0]
                if root_module in _FORBIDDEN_TOOL_MODULES:
                    return False, f"System Error: Rejected — forbidden module import: `{node.module}`.", set()
                if root_module in _ALLOWED_EXTERNAL_TOOL_MODULES:
                    detected_capabilities.add(root_module)
                elif root_module not in _SAFE_TOOL_MODULES:
                    return False, f"System Error: Rejected — unapproved module import: `{node.module}`.", set()

        # 2. Inspect identifiers (prevents calls, aliasing x=eval, capturing in structures [exec], passing as arguments)
        elif isinstance(node, ast.Name):
            if node.id in _FORBIDDEN_TOOL_CALLS or node.id in _FORBIDDEN_TOOL_DUNDERS:
                return False, f"System Error: Rejected — forbidden identifier: `{node.id}`.", set()
            if node.id in _FORBIDDEN_TOOL_MODULES:
                return False, f"System Error: Rejected — forbidden module reference: `{node.id}`.", set()

        # 3. Inspect attribute access (prevents dunder chaining and forbidden execution methods)
        elif isinstance(node, ast.Attribute):
            if node.attr in _FORBIDDEN_TOOL_DUNDERS:
                return False, f"System Error: Rejected — forbidden dunder attribute access: `{node.attr}`.", set()
            if node.attr in {"system", "popen", "spawn", "execv", "execve", "eval", "exec", "compile", "__import__"}:
                return False, f"System Error: Rejected — forbidden execution method: `{node.attr}()`.", set()
            if node.attr in {"sys", "os", "subprocess", "importlib", "builtins"}:
                return False, f"System Error: Rejected — forbidden module access: `{node.attr}`.", set()

    # 4. Verify exactly one top-level function with @tool decorator
    top_level_functions = [
        node for node in tree.body if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
    ]
    matching_functions = [node for node in top_level_functions if node.name == tool_name]
    if len(matching_functions) != 1:
        return (
            False,
            f"System Error: code must contain exactly one top-level function named '{tool_name}'.",
            set(),
        )

    def _decorator_name(dec):
        if isinstance(dec, ast.Call):
            return _decorator_name(dec.func)
        if isinstance(dec, ast.Name):
            return dec.id
        if isinstance(dec, ast.Attribute):
            base = _decorator_name(dec.value)
            return f"{base}.{dec.attr}" if base else dec.attr
        return ""

    def _has_tool_decorator(fn_node):
        return any(_decorator_name(dec).split(".")[-1] == "tool" for dec in fn_node.decorator_list)

    target_function = matching_functions[0]
    if not _has_tool_decorator(target_function):
        return False, f"System Error: function '{tool_name}' must have the @tool decorator.", set()

    extra_tool_functions = [
        node.name for node in top_level_functions
        if node.name != tool_name and _has_tool_decorator(node)
    ]
    if extra_tool_functions:
        return (
            False,
            f"System Error: only one @tool function is allowed. Extra decorated functions: {', '.join(extra_tool_functions)}.",
            set(),
        )

    return True, "", detected_capabilities


@tool
def write_custom_tool(tool_name: str, tool_code: str) -> str:
    """Writes and tests a new tool in astakos_skills/.
    It does not register it automatically in system/risk/registry — this is done with register_tool."""
    clean_code = re.sub(r"```(?:python)?", "", tool_code).replace("```", "").strip()

    valid, err_msg, detected_caps = _validate_custom_tool_ast(clean_code, tool_name)
    if not valid:
        return err_msg

    try:
        workspace_dir = os.path.realpath(WORKSPACE_DIR)
        final_path = os.path.realpath(os.path.join(workspace_dir, f"{tool_name}.py"))
        if not final_path.startswith(workspace_dir + os.sep):
            return "System Error: invalid tool path."
        if os.path.exists(final_path):
            return f"System Error: astakos_skills/{tool_name}.py already exists."
    except Exception:
        final_path = f"{tool_name}.py"

    test_script = f"""import math, json, inspect
from langchain_core.tools import tool

{clean_code}

if __name__ == "__main__":
    sig = inspect.signature({tool_name}.func if hasattr({tool_name}, 'func') else {tool_name})
    dummy = {{}}
    for p, param in sig.parameters.items():
        ann = param.annotation
        if ann in (int, float):
            dummy[p] = 1.0
        else:
            dummy[p] = "test"
    try:
        target_func = {tool_name}.func if hasattr({tool_name}, 'func') else {tool_name}
        if hasattr({tool_name}, 'invoke'):
            result = {tool_name}.invoke(dummy)
        else:
            result = target_func(**dummy)
        print(f"TEST_OK: {{result}}")
    except Exception as e:
        print(f"TEST_FAIL: {{e}}")
"""

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as tf:
            tf.write(test_script)
            temp_path = tf.name

        res = subprocess.run([sys.executable, temp_path], capture_output=True, text=True, timeout=15)
        stdout = res.stdout.strip()
        stderr = res.stderr.strip()

        if "TEST_FAIL" in stdout or (res.returncode != 0 and not stdout):
            error_detail = stdout or stderr
            return f"❌ Tool '{tool_name}' FAILED the test.\nError: {error_detail[:600]}"

        sep = "═" * 62
        code_body = re.sub(
            r"^\s*from\s+langchain_core\.tools\s+import\s+tool\s*\n",
            "",
            clean_code,
            flags=re.MULTILINE,
        ).strip()
        paste_code = f"from langchain_core.tools import tool\nimport math\n\n{code_body}"
        with open(final_path, "w", encoding="utf-8") as f:
            f.write(paste_code.rstrip() + "\n")

        caps_str = f" [Capabilities: {', '.join(sorted(detected_caps))}]" if detected_caps else ""
        print(f"\n\033[92m{sep}")
        print(f"  ✅  TOOL WRITTEN: {tool_name}{caps_str}")
        print(f"  🧪  Test: {stdout}")
        print(sep)
        print(paste_code)
        print(f"{sep}\033[0m\n")
        print(f"{config.USER_NAME}: ", end="", flush=True)

        return f"✅ Tool '{tool_name}' written to astakos_skills/{tool_name}.py and passed the test ({stdout}).{caps_str}"

    except subprocess.TimeoutExpired:
        return "❌ Timeout: the test script hung for more than 15 seconds."
    except Exception as e:
        return f"Error: {str(e)}"
    finally:
        if temp_path:
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception:
                pass


# ────────────────────────────────────────────────────────────────
# EMAIL
# ────────────────────────────────────────────────────────────────
import config
TOKEN_PATH = config.TOKEN_PATH
CREDS_PATH = config.CREDENTIALS_PATH

SCOPES = [
    'https://www.googleapis.com/auth/gmail.modify',
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/calendar',
    'https://www.googleapis.com/auth/tasks',
    'https://www.googleapis.com/auth/fitness.activity.read',
    'https://www.googleapis.com/auth/fitness.sleep.read',
    'https://www.googleapis.com/auth/fitness.heart_rate.read',
]

_RECOVERABLE_GOOGLE_OAUTH_REFRESH_ERRORS = (
    "invalid_scope",
    "invalid_grant",
)

def get_gmail_service():
    """Creates the Gmail API service using OAuth."""
    creds = None
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)

    if not creds or not creds.valid:
        if not os.path.exists(CREDS_PATH):
            raise Exception("Missing credentials.json! Download it from Google Cloud.")

        if creds and creds.expired and creds.refresh_token:
            from google.auth.transport.requests import Request
            try:
                creds.refresh(Request())
            except RefreshError as e:
                refresh_error = str(e).lower()
                if not any(
                    marker in refresh_error
                    for marker in _RECOVERABLE_GOOGLE_OAUTH_REFRESH_ERRORS
                ):
                    raise
                print("[GoogleAuth] token refresh rejected - forcing fresh OAuth consent.")
                creds = None

        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_secrets_file(CREDS_PATH, SCOPES)
            creds = flow.run_local_server(port=0, prompt='consent', access_type='offline')

        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())

    return build('gmail', 'v1', credentials=creds)


def decode_base64(data):
    return base64.urlsafe_b64decode(data.encode("UTF-8")).decode("utf-8", errors="replace")


def extract_body(payload):
    """Extracts body with fallback: plain text → HTML → nested parts."""
    import html

    def _parse_part(part):
        mime = part.get('mimeType', '')
        body = part.get('body', {})

        if 'parts' in part:
            for p in part['parts']:
                if p.get('mimeType') == 'text/plain' and 'data' in p.get('body', {}):
                    return decode_base64(p['body']['data'])
            for p in part['parts']:
                if p.get('mimeType') == 'text/html' and 'data' in p.get('body', {}):
                    return _html_to_text(decode_base64(p['body']['data']))
            for p in part['parts']:
                result = _parse_part(p)
                if result:
                    return result

        if 'data' in body:
            if mime == 'text/plain':
                return decode_base64(body['data'])
            elif mime == 'text/html':
                return _html_to_text(decode_base64(body['data']))

        return ""

    def _html_to_text(raw_html):
        # <br> and </p> → newlines for readable text
        raw_html = re.sub(r'<br\s*/?>', '\n', raw_html, flags=re.IGNORECASE)
        raw_html = re.sub(r'</p>', '\n', raw_html, flags=re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', raw_html)
        text = html.unescape(text)
        return re.sub(r'\s+', ' ', text).strip()

    return _parse_part(payload)


def clean_text(text):
    """Cleans whitespace and removes quoted replies (lines with >)."""
    lines = text.splitlines()
    lines = [l for l in lines if not l.strip().startswith('>')]
    cleaned = '\n'.join(lines)
    return re.sub(r'\s+', ' ', cleaned).strip()


def _encode_gmail_message(message: EmailMessage) -> str:
    return base64.urlsafe_b64encode(message.as_bytes()).decode("utf-8")


def _build_plain_email(to_email: str, subject: str, body: str) -> EmailMessage:
    message = EmailMessage()
    message["To"] = to_email
    message["Subject"] = subject
    message.set_content(body, charset="utf-8")
    return message


def _build_plain_reply(to_email: str, subject: str, body: str,
                       message_id: str = "", references: str = "") -> EmailMessage:
    message = _build_plain_email(to_email, subject, body)
    if message_id:
        message["In-Reply-To"] = message_id
    if references:
        message["References"] = references
    return message


@tool
def mail_manager(action: str, query: str = None, email_id: str = None,
                 to_email: str = None, subject: str = None, body: str = None,
                 limit: int = 10) -> str:
    """
    Gmail management via Google API. 
    Actions: 'search' (requires query), 'read_full' (requires email_id), 
             'read_thread' (requires email_id, reads the entire conversation),
             'send' (requires to_email, subject, body),
             'reply' (requires email_id, body),
             'delete' (requires email_id).
    """
    try:
        if not action:
            return "❌ Provide action: search, read_full, read_thread, send, reply or delete."

        print(f"\033[94m[Mail API]: Executing action '{action}'...\033[0m")
        action = action.lower()
        if action == "read" and email_id:
            action = "read_full"
        elif action == "read":
            action = "search"
        service = get_gmail_service()

        # =========================
        # SEND
        # =========================
        if action == "send":
            if not to_email or not subject or not body:
                return "❌ Send requires: to_email, subject, body."
            raw = _encode_gmail_message(_build_plain_email(to_email, subject, body))
            service.users().messages().send(userId="me", body={"raw": raw}).execute()
            return "✅ Email sent successfully."

        # =========================
        # REPLY
        # =========================
        elif action == "reply":
            if not email_id or not body:
                return "❌ Reply requires email_id and body."
            
            original = service.users().messages().get(
                userId="me", id=email_id, format="metadata",
                metadataHeaders=["Subject", "From", "Message-ID", "References"]
            ).execute()
            
            headers = original["payload"]["headers"]
            orig_subject = next((h["value"] for h in headers if h["name"] == "Subject"), "Re: (no subject)")
            orig_from    = next((h["value"] for h in headers if h["name"] == "From"), "")
            orig_msg_id  = next((h["value"] for h in headers if h["name"] == "Message-ID"), "")
            orig_refs    = next((h["value"] for h in headers if h["name"] == "References"), "")
            
            reply_subject = orig_subject if orig_subject.startswith("Re:") else f"Re: {orig_subject}"
            references = f"{orig_refs} {orig_msg_id}".strip()

            reply_message = _build_plain_reply(
                orig_from,
                reply_subject,
                body,
                message_id=orig_msg_id,
                references=references,
            )
            raw = _encode_gmail_message(reply_message)
            thread_id = original.get("threadId")
            send_body = {"raw": raw}
            if thread_id:
                send_body["threadId"] = thread_id
            service.users().messages().send(
                userId="me",
                body=send_body
            ).execute()
            return f"✅ Reply sent to {orig_from}."

        # =========================
        # SEARCH
        # =========================
        elif action in ["search", "check_emails", "check"]:
            results = service.users().messages().list(userId="me", q=query, maxResults=limit).execute()
            messages = results.get("messages", [])

            if not messages:
                return f"No emails found for search: {query}"

            output = []
            for msg in messages:
                data = service.users().messages().get(
                    userId="me", id=msg['id'],
                    format="metadata",
                    metadataHeaders=["Subject", "From", "Date"]
                ).execute()
                headers = data['payload']['headers']
                subject_val = next((h['value'] for h in headers if h['name'] == 'Subject'), "No Subject")
                from_val    = next((h['value'] for h in headers if h['name'] == 'From'), "Unknown")
                date_val    = next((h['value'] for h in headers if h['name'] == 'Date'), "")
                output.append(f"ID: {msg['id']} | {date_val} | From: {from_val} | Subject: {subject_val}")

            return "\n".join(output)

        # =========================
        # READ FULL (Single Message)
        # =========================
        elif action == "read_full":
            if not email_id:
                return "❌ Read_full requires email_id."
            data = service.users().messages().get(userId="me", id=email_id, format="full").execute()
            body_text = extract_body(data['payload'])
            return f"{t('prompts.mail_content_result_prefix')}\n{clean_text(body_text)[:5000]}"

        # =========================
        # READ THREAD (Full Conversation)
        # =========================
        elif action == "read_thread":
            if not email_id:
                return "❌ Read_thread requires email_id (of a message in the thread)."
            # Get the message to find its threadId
            msg_meta = service.users().messages().get(userId="me", id=email_id, format="minimal").execute()
            thread_id = msg_meta.get("threadId", email_id)
            
            thread_data = service.users().threads().get(userId="me", id=thread_id).execute()
            messages_in_thread = thread_data.get("messages", [])
            
            output = []
            for i, m in enumerate(messages_in_thread):
                headers = m['payload']['headers']
                from_val = next((h['value'] for h in headers if h['name'] == 'From'), "Unknown")
                date_val = next((h['value'] for h in headers if h['name'] == 'Date'), "")
                
                body_text = extract_body(m['payload'])
                output.append(f"--- Message {i+1} | From: {from_val} ({date_val}) ---\n{clean_text(body_text)[:2000]}")
            
            full_text = "\n\n".join(output)
            return (
                f"{t('prompts.mail_thread_result_prefix')} "
                f"({len(messages_in_thread)} messages):\n{full_text[:8000]}"
            )

        # =========================
        # DELETE
        # =========================
        elif action == "delete":
            if not email_id:
                return "❌ Delete requires email_id."
            service.users().messages().trash(userId="me", id=email_id).execute()
            return f"🗑️ Email {email_id} moved to trash."

        return f"❌ Unknown command: {action}"

    except Exception as e:
        return f"Mail API Error: {str(e)}"

# ────────────────────────────────────────────────────────────────
# GITHUB
# ────────────────────────────────────────────────────────────────

import subprocess
import shlex
from langchain_core.tools import tool

@tool
def github_manager(action: str, repo_name: str = "", target_files: str = "",
                   commit_message: str = "", content: str = "") -> str:
    """
    Manages GitHub operations and local Git commits.
    
    Actions: 
    - 'list_repos', 'read_file', 'create_file', 'update_file' (Cloud API Operations)
    - 'push_local_commits' (Runs local Git CLI commands)
    
    CRITICAL RULES for 'push_local_commits': 
    - target_files MUST be a comma-separated list of exact paths (e.g. "core/brain.py, api/server.py"). 
    - Using "." or "*" or "all" is STRICTLY FORBIDDEN.
    """
    token = os.getenv("GITHUB_TOKEN") 
    if not token:
        return "Error: Missing GITHUB_TOKEN."

    try:
        from github import Github
        # ─── 1. LOCAL GIT CLI OPERATIONS (Mastro-Shielded) ──────────────
        if action == "push_local_commits":
            if not target_files or target_files.strip() in [".", "*", "all"]:
                return "🛡️ [GIT OVERRIDE]: Blind sweeps are forbidden. Specify exact file paths."
            if not commit_message:
                return "Error: Commit message is required."

            # ── SafeExec check ───────────────────────────────────────────
            from core.safe_executor import safe_execute
            push_check = safe_execute("git push origin main", lambda c: {"status": "ok"})
            if push_check.get("status") == "blocked":
                return f"🛡️ [SAFE EXECUTOR - BLOCKED]: {push_check['reason']}"
            if push_check.get("status") == "cancelled":
                return "⚠️ [SAFE EXECUTOR]: Git push requires confirmation. Send again with `/confirm`"
            # ────────────────────────────────────────────────────────────

            files = [f.strip() for f in target_files.split(",") if f.strip()]

            # 1. git add <files>
            add_cmd = ["git", "add"] + files
            subprocess.run(add_cmd, check=True, capture_output=True, text=True)

            # 2. git commit -m "message"
            commit_cmd = ["git", "commit", "-m", commit_message]
            subprocess.run(commit_cmd, check=True, capture_output=True, text=True)

            # 3. git push origin main
            push_cmd = ["git", "push", "origin", "main"]
            subprocess.run(push_cmd, check=True, capture_output=True, text=True)

            return f"System: Local changes successfully pushed!\nFiles: {files}\nMessage: {commit_message}"

        # ─── 2. GITHUB CLOUD API OPERATIONS ─────────────────────────────
        g = Github(token)
        user = g.get_user()

        if action == "list_repos":
            repos = [f"- {r.name} ({'Private' if r.private else 'Public'})" for r in user.get_repos()]
            return f"Found {len(repos)} Repositories:\n" + "\n".join(repos)

        elif action == "read_file":
            repo = g.get_repo(f"{user.login}/{repo_name}")
            file_content = repo.get_contents(target_files)
            return f"Content of {target_files}:\n{file_content.decoded_content.decode('utf-8')[:10000]}"

        elif action in ["create_file", "update_file"]:
            if not content.strip():
                return "🛡️ [GIT OVERRIDE]: Content is empty. Refusing to overwrite file with empty data."
            repo = g.get_repo(f"{user.login}/{repo_name}")
            try:
                file_info = repo.get_contents(target_files)
                repo.update_file(target_files, commit_message, content, file_info.sha)
                return f"System: '{target_files}' in '{repo_name}' updated via API!"
            except:
                repo.create_file(target_files, commit_message, content)
                return f"System: '{target_files}' created in '{repo_name}' via API!"

        else:
            return "Error: Invalid action specified."

    except subprocess.CalledProcessError as e:
        return f"Local Git Command Error: {e.stderr}"
    except Exception as e:
        return f"GitHub API Error: {str(e)}"


# ────────────────────────────────────────────────────────────────
# HARDWARE CONTROL
# ────────────────────────────────────────────────────────────────

@tool
def control_vacuum(action: str) -> str:
    """Controls the Xiaomi X20+ robot vacuum.
    Actions: 'start', 'stop', 'home', 'status' (for battery/status), 'room:<room_name>' (e.g., room:Kitchen)."""
    ip = VACUUM_IP
    token = VACUUM_TOKEN

    if not ip or not token:
        return "Error: VACUUM_IP or VACUUM_TOKEN not found."

    try:
        from miio import Device
        vac = Device(ip, token)

        if action == "start":
            vac.send("action", {"did": "astakos", "siid": 2, "aiid": 1, "in": []})
            return t("tools.system.msg_vacuum_start", bot_name=config.BOT_NAME)

        elif action == "stop":
            vac.send("action", {"did": "astakos", "siid": 2, "aiid": 2, "in": []})
            return t("tools.system.msg_vacuum_stop", bot_name=config.BOT_NAME)

        elif action == "home":
            vac.send("action", {"did": "astakos", "siid": 3, "aiid": 1, "in": []})
            return t("tools.system.msg_vacuum_return", bot_name=config.BOT_NAME)

        elif action.startswith("room:"):
            room_name = action.split(":", 1)[1].strip()
            import os, json
            
            # Loading of room_map.json
            room_map_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "room_map.json")
            try:
                with open(room_map_path, "r", encoding="utf-8") as f:
                    room_map = json.load(f)
            except Exception:
                # Fallback if the file does not exist (we have hardcoded the ones we found)
                room_map = t("tools.system.room_map_fallback")

            room_id = room_map.get(room_name)
            if room_id is None:
                available = ", ".join(room_map.keys())
                return f"Error: Not found: room '{room_name}'. Available rooms: {available}"

            vac.send("action", {
                "did": "astakos", 
                "siid": 4, 
                "aiid": 1, 
                "in": [
                    {"piid": 1, "value": 18}, 
                    {"piid": 10, "value": f'{{"selects":[[{room_id},1,2,1,1]]}}'}
                ]
            })
            return t("tools.system.msg_vacuum_room", bot_name=config.BOT_NAME, room_name=room_name)

        else:
            return f"Unknown command: {action}."

    except Exception as e:
        return f"Error communicating with vacuum: {str(e)}"
@tool
def post_to_linkedin(text: str = None, image_path: str = None, image_paths: str = None) -> str:
    """
    Publishes text and optionally images to LinkedIn.
    If no text is provided, it is automatically retrieved from linkedin_draft.json.
    image_path: a single image (backward compatibility)
    image_paths: multiple images separated by commas, e.g., "C:/a.jpg,C:/b.jpg,C:/c.jpg"
                 The images are uploaded as a carousel on LinkedIn (up to 9).
    """
    import os
    import json
    import requests
    from dotenv import load_dotenv, find_dotenv
    from config import LINKEDIN_DRAFT_FILE

    # [MASTRO-INTERCEPTOR]: Autonomy from Working Memory
    if not text:
        if os.path.exists(LINKEDIN_DRAFT_FILE):
            try:
                with open(LINKEDIN_DRAFT_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    text = data.get("content") or data.get("text")
                    if not image_paths and not image_path:
                        image_paths = data.get("image_paths")
                        image_path = data.get("image_path")
            except Exception as e:
                print(f"⚠️ Error reading draft: {e}")

    if not text:
        return "❌ Error: No text found (neither in draft nor in arguments)."

    # Gathering of all paths into a list
    all_paths = []
    if image_paths:
        all_paths = [p.strip() for p in image_paths.split(",") if p.strip()]
    elif image_path:
        all_paths = [image_path]

    # Validate paths
    for p in all_paths:
        if not os.path.exists(p):
            return f"❌ Image not found: {p}"

    # --- LinkedIn API Logic ---
    load_dotenv(find_dotenv(), override=True)
    token = os.getenv("LINKEDIN_TOKEN")
    if not token: return "❌ Missing LINKEDIN_TOKEN."

    headers = {
        "Authorization": f"Bearer {token}",
        "X-Restli-Protocol-Version": "2.0.0",
        "Content-Type": "application/json"
    }

    try:
        # 1. Identification
        user_res = requests.get("https://api.linkedin.com/v2/userinfo", headers=headers)
        if user_res.status_code != 200: return f"❌ Auth Error: {user_res.text}"
        person_urn = f"urn:li:person:{user_res.json().get('sub')}"

        # 2. Upload Images (single or multiple)
        asset_urns = []
        for path in all_paths[:9]:  # LinkedIn max 9 images
            reg_url = "https://api.linkedin.com/v2/assets?action=registerUpload"
            reg_data = {"registerUploadRequest": {"recipes": ["urn:li:digitalmediaRecipe:feedshare-image"], "owner": person_urn, "serviceRelationships": [{"relationshipType": "OWNER", "identifier": "urn:li:userGeneratedContent"}]}}
            reg_res = requests.post(reg_url, headers=headers, json=reg_data)
            if reg_res.status_code == 200:
                upload_url = reg_res.json()['value']['uploadMechanism']['com.linkedin.digitalmedia.uploading.MediaUploadHttpRequest']['uploadUrl']
                asset_urn = reg_res.json()['value']['asset']
                with open(path, 'rb') as f:
                    requests.post(upload_url, headers={"Authorization": f"Bearer {token}"}, data=f.read())
                asset_urns.append(asset_urn)

        # 3. Create Post
        post_url = "https://api.linkedin.com/v2/ugcPosts"
        media_content = {
            "shareCommentary": {"text": text},
            "shareMediaCategory": "IMAGE" if asset_urns else "NONE"
        }
        if asset_urns:
            media_content["media"] = [
                {"status": "READY", "media": urn, "title": {"text": f"Photo {i+1}"}}
                for i, urn in enumerate(asset_urns)
            ]

        payload = {
            "author": person_urn,
            "lifecycleState": "PUBLISHED",
            "specificContent": {"com.linkedin.ugc.ShareContent": media_content},
            "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"}
        }

        res = requests.post(post_url, headers=headers, json=payload)

        if res.status_code == 201:
            # [MASTRO-CLEANUP]: Draft cleanup after success
            if os.path.exists(LINKEDIN_DRAFT_FILE):
                with open(LINKEDIN_DRAFT_FILE, "w", encoding="utf-8") as f:
                    json.dump({}, f)
            img_count = len(asset_urns)
            img_msg = f" with {img_count} image{'s' if img_count != 1 else ''}" if img_count else ""
            return f"✅ LinkedIn post uploaded{img_msg} and draft cleared!"

        return f"❌ Failure: {res.text}"

    except Exception as e:
        return f"❌ Critical Error: {str(e)}"
import math

def _is_home(lat: float, lon: float, home_lat: float = 0.0, home_lon: float = 0.0, radius_m: float = 150) -> bool:
    """Checks if the coordinates are within 150 meters of {config.DEVELOPER_NAME}."""
    R = 6371000
    dlat = math.radians(lat - home_lat)
    dlon = math.radians(lon - home_lon)
    a = math.sin(dlat/2)**2 + math.cos(math.radians(home_lat)) * math.cos(math.radians(lat)) * math.sin(dlon/2)**2
    return R * 2 * math.asin(math.sqrt(a)) < radius_m


@tool
def get_current_location() -> str:
    """
    Returns the last recorded GPS coordinate of {config.USER_NAME} from last_location.json.
    Used to know where the user is in real-time.
    """
    import json
    import os
    import time
    from datetime import datetime
    from config import GPS_STORAGE_FILE

    if not os.path.exists(GPS_STORAGE_FILE):
        return "📍 No recorded location found. Ask {config.USER_NAME} to send Live Location."

    try:
        with open(GPS_STORAGE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            lat = data.get("lat")
            lon = data.get("lon")
            ts = data.get("timestamp", 0)

            # Calculation of "freshness"
            diff_minutes = int((time.time() - ts) / 60)
            last_seen = datetime.fromtimestamp(ts).strftime('%H:%M:%S')

            if diff_minutes > 1440:
                return f"📍 Location is very old ({diff_minutes // 60}h old, last updated {last_seen})."

            maps_link = f"https://maps.google.com/?q={lat},{lon}"
            home_status = "🏠 Is HOME" if _is_home(float(lat), float(lon)) else "🚶 Is OUT of home"

            return (
                f"📍 Coordinates: {lat}, {lon}\n"
                f"{home_status}\n"
                f"🗺️ <a href='{maps_link}'>View on Map</a>\n"
                f"⏱️ Updated {diff_minutes} minutes ago (at {last_seen})."
            )

    except Exception as e:
        return f"❌ Error reading GPS: {str(e)}"

@tool
def control_spotify(
    action: str,
    query: str = ""
) -> str:
    """Controls Spotify.
    action: 'play', 'pause', 'next', 'now_playing', 'top_tracks', 'search'
    query: Title/Artist for action='search'"""
    try:
        scope = "user-modify-playback-state user-read-playback-state user-top-read user-read-currently-playing"
        sp = spotipy.Spotify(auth_manager=SpotifyOAuth(scope=scope))

        if action == "top_tracks":
            results = sp.current_user_top_tracks(limit=5, time_range='long_term')
            if not results['items']:
                return "No top tracks data found."
            tracks = [f"{i+1}. {t['name']} - {t['artists'][0]['name']}" for i, t in enumerate(results['items'])]
            return "🎵 Your Top 5 tracks:\n" + "\n".join(tracks)

        elif action == "pause":
            sp.pause_playback()
            return "⏸️ Music paused."

        elif action == "next":
            sp.next_track()
            return "⏭️ Skipped to next track!"

        elif action == "now_playing":
            current = sp.current_playback()
            if not current or not current.get("item"):
                return "Nothing playing right now."
            track = current["item"]
            artist = track["artists"][0]["name"]
            name = track["name"]
            playing = "▶️" if current["is_playing"] else "⏸️"
            return f"{playing} {name} — {artist}"

        elif action == "search":
            if not query:
                return "❌ Provide title or artist to search."
            res = sp.search(q=query, type='track', limit=1)
            if not res['tracks']['items']:
                return f"❌ Could not find '{query}'."
            track_uri = res['tracks']['items'][0]['uri']
            track_name = res['tracks']['items'][0]['name']
            sp.start_playback(uris=[track_uri])
            return f"▶️ Now playing: {track_name} 🎵"

        elif action == "play":
            sp.start_playback()
            return "▶️ Music resumed!"

        return "❌ Unknown command. Try: play, pause, next, now_playing, top_tracks, search."

    except Exception as e:
        return f"⚠️ Spotify Error: {str(e)}. (Is the app open?)"

@tool
def get_fit_summary(days_ago: int = 1) -> str:
    """
    Returns a Google Fit summary for {config.USER_NAME}.
    days_ago=0 → today, days_ago=1 → yesterday (default).
    Includes: steps, sleep (hours + deep/REM), heart rate.
    """
    try:
        from astakos_skills.google_fit import get_daily_summary
        return get_daily_summary(days_ago=days_ago)
    except Exception as e:
        return f"❌ Google Fit error: {e}"


@tool
def save_goal_tool(
    project: str,
    description: str,
    status: str = "active",
    progress: int = 0,
    milestones: str = "",
    external_content_sources_json: str = "",
) -> str:
    """
    Saves or updates a long-term goal for {config.USER_NAME}.
    project: Short project name (e.g., 'ShiftMaster', 'Astakos', 'PraxisERP').
    description: What he wants to achieve (e.g., 'To finish the licensing module').
    status: 'active' (in progress) | 'paused' (shelved) | 'done' (completed).
    progress: Progress percentage 0-100.
    milestones: Smaller steps or milestones (as a string).
    external_content_sources_json: Internal approval provenance. Do not set this manually.
    """
    from memory.vector_store import save_goal
    from core.untrusted_content import external_content_sources_from_json

    ok = save_goal(
        project=project,
        description=description,
        status=status,
        progress=progress,
        milestones=milestones,
        external_content_sources=external_content_sources_from_json(
            external_content_sources_json,
        ),
    )
    if ok:
        return f"✅ Goal '{project}' saved ({status}, {progress}%)."
    return f"❌ Failed to save goal '{project}'."


@tool
def update_goal_status_tool(project: str, status: str) -> str:
    """
    Updates the status of an existing goal.
    project: The name of the project (e.g., 'ShiftMaster').
    status: 'active' | 'paused' | 'done'
    """
    from memory.vector_store import update_goal_status
    ok = update_goal_status(project=project, status=status)
    if ok:
        return f"✅ Goal '{project}' → {status}."
    return f"❌ Goal '{project}' not found."


@tool
def update_goal_progress_tool(project: str, progress: int) -> str:
    """
    Updates the progress percentage of an existing goal (0-100).
    project: The name of the project.
    progress: An integer from 0 to 100.
    """
    from memory.vector_store import update_goal_progress
    ok = update_goal_progress(project=project, progress=progress)
    if ok:
        return f"✅ Goal '{project}' progress → {progress}%."
    return f"❌ Goal '{project}' not found."


@tool
def update_goal_milestones_tool(
    project: str,
    milestones: str,
    external_content_sources_json: str = "",
) -> str:
    """
    Updates the milestones of an existing goal.
    project: The name of the project.
    milestones: The new milestones (in string format, e.g., '1) UI, 2) DB').
    """
    from memory.vector_store import update_goal_milestones
    from core.untrusted_content import external_content_sources_from_json

    ok = update_goal_milestones(
        project=project,
        milestones=milestones,
        external_content_sources=external_content_sources_from_json(
            external_content_sources_json,
        ),
    )
    if ok:
        return f"✅ Goal '{project}' milestones updated."
    return f"❌ Goal '{project}' not found."


@tool
def tool_stats(days: int = 7) -> str:
    """
    Displays performance statistics of the tools for the last N days.
    days: How many days back to look (default 7).
    Returns per tool: calls, errors, error rate, average duration.
    """
    import os, json
    from datetime import datetime, timedelta
    from collections import defaultdict

    traces_dir = os.path.join(os.path.dirname(__file__), "..", "logs", "traces")
    if not os.path.isdir(traces_dir):
        return "❌ Traces folder not found."

    stats: dict[str, dict] = defaultdict(lambda: {"calls": 0, "errors": 0, "durations": []})

    today = datetime.now().date()
    loaded_days = 0
    for i in range(days):
        day = today - timedelta(days=i)
        path = os.path.join(traces_dir, f"{day}.json")
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                entries = json.load(f)
        except Exception:
            continue
        loaded_days += 1
        for entry in entries:
            for tc in entry.get("tool_calls", []):
                name = tc.get("tool", "unknown")
                stats[name]["calls"] += 1
                if tc.get("error"):
                    stats[name]["errors"] += 1
                dur = tc.get("duration_ms")
                if dur is not None:
                    stats[name]["durations"].append(dur)

    if not stats:
        return f"📊 No traces found for the last {days} days."

    # Sorting: first those with errors, then alphabetically
    rows = []
    for name, s in sorted(stats.items(), key=lambda x: (-x[1]["errors"], x[0])):
        calls = s["calls"]
        errors = s["errors"]
        rate = f"{errors/calls*100:.0f}%" if calls else "—"
        avg_dur = f"{sum(s['durations'])//len(s['durations'])}ms" if s["durations"] else "—"
        err_icon = "🔴" if errors > 0 else "✅"
        rows.append(f"{err_icon} {name}: {calls} calls, {errors} errors ({rate}), avg {avg_dur}")

    header = f"📊 Tool Stats — last {days} days ({loaded_days} trace files)\n"
    return header + "\n".join(rows)


def _doctor_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _doctor_load_json(path: str):
    try:
        if not os.path.exists(path):
            return []
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _doctor_event_has_error(event: dict) -> bool:
    action = str(event.get("action", "")).lower()
    return bool(event.get("error")) or action in {"error", "failed", "failure", "exception"} or "error" in action


def _doctor_trace_has_issue(trace: dict) -> bool:
    if trace.get("error") or trace.get("loop_guard"):
        return True
    return any(tc.get("error") for tc in trace.get("tool_calls", []))


def _doctor_summarize_logs(days: int = 2, *, root: str | None = None, slow_ms: int = 45000) -> dict:
    root = root or _doctor_root()
    days = max(1, int(days or 1))
    today = datetime.now().date()
    summary = {
        "days": days,
        "event_files": 0,
        "events": 0,
        "event_errors": 0,
        "trace_files": 0,
        "traces": 0,
        "trace_issues": 0,
        "loop_guards": 0,
        "slow_traces": 0,
        "last_issues": [],
    }

    for i in range(days):
        day = today - timedelta(days=i)
        day_str = day.isoformat()

        events_path = os.path.join(root, "logs", "events", f"{day_str}.json")
        events = _doctor_load_json(events_path)
        if events:
            summary["event_files"] += 1
        summary["events"] += len(events)
        summary["event_errors"] += sum(1 for event in events if _doctor_event_has_error(event))

        traces_path = os.path.join(root, "logs", "traces", f"{day_str}.json")
        traces = _doctor_load_json(traces_path)
        if traces:
            summary["trace_files"] += 1
        summary["traces"] += len(traces)
        for trace in traces:
            is_issue = _doctor_trace_has_issue(trace)
            is_loop = bool(trace.get("loop_guard"))
            is_slow = int(trace.get("duration_ms") or 0) >= slow_ms
            summary["trace_issues"] += int(is_issue)
            summary["loop_guards"] += int(is_loop)
            summary["slow_traces"] += int(is_slow)
            if (is_issue or is_slow) and len(summary["last_issues"]) < 5:
                label = "loop guard" if is_loop else "error" if is_issue else "slow"
                summary["last_issues"].append({
                    "timestamp": trace.get("timestamp", "?"),
                    "agent": trace.get("agent") or "?",
                    "label": label,
                    "message": str(trace.get("user_message") or "")[:90],
                })

    return summary


def _doctor_compact_map(values: dict | None) -> str:
    if not values:
        return "none"
    return ", ".join(f"{k}:{v}" for k, v in values.items())


def _count_memory_audit_ops(entries: list[dict]) -> dict[str, int]:
    counts = {
        "add": 0,
        "overwrite": 0,
        "add_alongside": 0,
        "skip_duplicate": 0,
        "skip_keep_old": 0,
        "reflection": 0,
        "total": len(entries),
    }
    for entry in entries:
        op = entry.get("op")
        if op in counts:
            counts[op] += 1
        elif op in ("reflection_applied", "reflection_saved"):
            counts["reflection"] += 1
    return counts


def _format_memory_ops_summary(counts: dict[str, int]) -> str:
    if not counts.get("total"):
        return "0 ops"
    parts = [
        f"add {counts.get('add', 0)}",
        f"overwrite {counts.get('overwrite', 0)}",
        f"alongside {counts.get('add_alongside', 0)}",
        f"skipped {counts.get('skip_duplicate', 0) + counts.get('skip_keep_old', 0)}",
        f"reflections {counts.get('reflection', 0)}",
    ]
    return f"{counts.get('total', 0)} ops (" + ", ".join(parts) + ")"


def _format_pending_routines(pending_routines: dict) -> str:
    if not pending_routines:
        return "0"
    names = []
    for routine_id, item in list(pending_routines.items())[:3]:
        event = item.get("event") if isinstance(item, dict) else ""
        event = str(event or f"routine #{routine_id}").strip()
        names.append(event[:60])
    suffix = ", ".join(names)
    if len(pending_routines) > 3:
        suffix += f", +{len(pending_routines) - 3}"
    return f"{len(pending_routines)} — {suffix}"


def _doctor_status_label(*, warnings: list[str], pending_actions: list, logs: dict) -> str:
    if logs.get("event_errors") or any("unreadable" in w for w in warnings):
        return "Immediate check"
    if pending_actions or logs.get("trace_issues") or logs.get("loop_guards") or warnings:
        return "Warning"
    return "OK"


@tool
def system_doctor(days: int = 1) -> str:
    """
    Fast read-only health check of the Astakos runtime.
    days: how many days of logs/traces to look at (default 1, i.e. today).
    """
    lines: list[str] = []
    warnings: list[str] = []

    logs = _doctor_summarize_logs(days=days)
    if logs["event_errors"]:
        warnings.append(f"{logs['event_errors']} event errors")
    if logs["trace_issues"]:
        warnings.append(f"{logs['trace_issues']} trace issues")
    if logs["loop_guards"]:
        warnings.append(f"{logs['loop_guards']} loop guards")

    try:
        from core.approval import list_pending
        pending_actions = list_pending()
    except Exception:
        pending_actions = []
        warnings.append("approval store unreadable")

    try:
        from core.messenger_draft import debug_draft_state
        draft = debug_draft_state()
    except Exception:
        draft = {"exists": False, "active": False}
        warnings.append("draft state unreadable")

    try:
        from memory.conversation_history import load_conversation_stats
        from memory.session_memory import AUTO_SESSION_SUMMARY_EXCHANGE_THRESHOLD
        conv = load_conversation_stats()
        unsummarized = int(conv.get("unsummarized_exchanges") or 0)
        threshold = int(AUTO_SESSION_SUMMARY_EXCHANGE_THRESHOLD)
        if unsummarized >= threshold:
            warnings.append("session summary due")
    except Exception:
        conv = {}
        unsummarized = 0
        threshold = 0
        warnings.append("conversation stats unreadable")

    try:
        from memory.routine_db import load_pending_confirmations
        pending_routines = load_pending_confirmations()
    except Exception:
        pending_routines = {}
        warnings.append("routine pending store unreadable")

    try:
        memory_ops = _count_memory_audit_ops(_load_audit_log(days=days))
    except Exception:
        memory_ops = {"total": 0}
        warnings.append("memory audit unreadable")

    try:
        cond_routines = _doctor_conditioned_routines()
    except Exception:
        cond_routines = "error reading conditions"

    status = _doctor_status_label(warnings=warnings, pending_actions=pending_actions, logs=logs)
    lines.append(t("tools.system.msg_doctor_status", bot_name=config.BOT_NAME, status=status))
    lines.append(t("tools.system.msg_doctor_logs", days=logs['days'], events=logs['events'], event_errors=logs['event_errors'], traces=logs['traces'], trace_issues=logs['trace_issues']))
    lines.append(t("tools.system.msg_doctor_loops", loop_guards=logs['loop_guards'], slow_traces=logs['slow_traces']))
    lines.append(t("tools.system.msg_doctor_approvals", count=len(pending_actions), tools=", ".join(a.get('tool_name', '?') for a in pending_actions[:3]) if pending_actions else ""))
    lines.append(t("tools.system.msg_doctor_draft", active=draft.get('active'), target=draft.get('target_name') if draft.get("active") and draft.get("target_name") else ""))
    lines.append(t("tools.system.msg_doctor_backlog", unsummarized=unsummarized, threshold=threshold, channels=_doctor_compact_map(conv.get('unsummarized_by_channel'))))
    lines.append(f"• Memory ops: {_format_memory_ops_summary(memory_ops)}")
    lines.append(f"• Pending routine confirmations: {_format_pending_routines(pending_routines)}")

    if cond_routines and cond_routines != "None":
        lines.append("• Conditioned routines:")
        lines.append(cond_routines)

    try:
        ctx_panel = _doctor_runtime_context()
        if ctx_panel and ctx_panel != "None":
            lines.append("• Runtime Context:")
            for c_line in ctx_panel.splitlines():
                lines.append(f"  {c_line}")
    except Exception:
        pass

    if logs["last_issues"]:
        lines.append("• Recent things to inspect:")
        for item in logs["last_issues"][-3:]:
            lines.append(f"  - {item['timestamp']} [{item['label']}] {item['agent']}: {item['message']}")

    if warnings:
        lines.append("• Note: " + ", ".join(warnings[:5]))
    else:
        lines.append("• Everything looks quiet.")

    return "\n".join(lines)


def _doctor_runtime_context() -> str:
    try:
        from services.routine_context import build_runtime_routine_context
        ctx = build_runtime_routine_context()
        if not ctx:
            return "None"
        import json
        return json.dumps(ctx, indent=2, ensure_ascii=False)
    except Exception as e:
        return f"Error: {e}"

def _doctor_conditioned_routines() -> str:
    try:
        from memory.routine_db import get_connection
        from services.routine_conditions import evaluate_routine_condition
        from services.routine_context import build_runtime_routine_context

        conn = get_connection()
        cursor = conn.cursor()
        rows = cursor.execute(
            """
            SELECT id, event_name, condition_type, condition_mode, condition_payload, priority, conflict_group
            FROM routines
            WHERE state = 'active' AND (condition_type IS NOT NULL OR priority > 0 OR conflict_group IS NOT NULL)
            """
        ).fetchall()
        conn.close()

        if not rows:
            return "None"

        context = build_runtime_routine_context()
        lines = []
        for r_id, event_name, c_type, c_mode, c_payload, priority, conflict_group in rows:
            eval_result = {}
            if c_type:
                eval_result = evaluate_routine_condition(
                    {
                        "condition_type": c_type,
                        "condition_payload": c_payload,
                        "condition_mode": c_mode,
                    },
                    context,
                )
            
            status = "allowed" if eval_result.get("allowed", True) else "blocked"
            eval_reason = eval_result.get("reason", "")
            
            mode_desc = f"{c_mode}" if c_mode else "no_condition"
            
            import json
            try:
                payload_dict = json.loads(c_payload) if c_payload else {}
                if "flag" in payload_dict:
                    target = f"{payload_dict['flag']}={payload_dict.get('equals', True)}"
                elif "shift" in payload_dict:
                    target = f"shift={payload_dict['shift']}"
                else:
                    target = str(c_payload)
            except Exception:
                target = str(c_payload)

            details = []
            if c_type:
                details.append(f"cond: {mode_desc} ({target}) -> {status} [{eval_reason}]")
            if priority or conflict_group:
                details.append(f"conflict: grp='{conflict_group or event_name.split()[0]}' prio={priority or 0}")
                
            lines.append(f"  #{r_id} {event_name} | " + " | ".join(details))

        if not lines:
            return "None"
        return "\n".join(lines)
    except Exception as e:
        return f"error: {str(e)}"

# ────────────────────────────────────────────────────────────────
# MEMORY REVIEW
# ────────────────────────────────────────────────────────────────

def _load_audit_log(days: int = 1) -> list[dict]:
    """Loads entries from the memory audit log for the last X days."""
    from config import MEMORY_AUDIT_DIR
    from datetime import date, timedelta
    entries = []
    today = date.today()
    for i in range(days):
        day = today - timedelta(days=i)
        path = os.path.join(MEMORY_AUDIT_DIR, f"{day.isoformat()}.json")
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        entries.extend(data)
            except Exception:
                pass
    return entries


def _memory_review_period(days: int) -> str:
    return "Today" if days <= 1 else f"In the last {days} days"


def _normalize_memory_review_op(op: str | None) -> str:
    value = (op or "").strip().lower()
    aliases = NLP_CONFIG.get("system", {}).get("memory_review_aliases", {})
    return aliases.get(value, value)


def _filter_memory_audit_entries(entries: list[dict], *, op: str = "", category: str = "") -> list[dict]:
    normalized_op = _normalize_memory_review_op(op)
    category_value = (category or "").strip().lower()
    filtered = entries
    if normalized_op:
        if normalized_op == "skip":
            filtered = [e for e in filtered if e.get("op") in ("skip_duplicate", "skip_keep_old")]
        elif normalized_op == "reflection":
            filtered = [e for e in filtered if e.get("op") in ("reflection_applied", "reflection_saved")]
        else:
            filtered = [e for e in filtered if e.get("op") == normalized_op]
    if category_value:
        filtered = [e for e in filtered if str(e.get("category", "")).strip().lower() == category_value]
    return filtered


def _memory_review_has_filter(op: str = "", category: str = "") -> bool:
    return bool((op or "").strip() or (category or "").strip())


def _append_memory_review_section(
    lines: list[str],
    *,
    title: str,
    entries: list[dict],
    has_filter: bool,
    limit: int,
    formatter,
) -> None:
    if has_filter and not entries:
        return
    lines.append(f"\n{title}: {len(entries)}")
    for entry in entries[-limit:]:
        lines.append(formatter(entry))


@tool
def memory_review(days: int = 1, op: str = "", category: str = "") -> str:
    """
    Displays what Astakos stored in memory during the last X days.
    Includes: new entries, overwrites, duplicates skipped, reflections.
    days: how many days back (default=1 = today).
    op: optional filter (add, overwrite, add_alongside, skip, skip_duplicate, skip_keep_old, reflection).
    category: optional category filter (e.g., family, lazeros, projects).
    """
    all_entries = _load_audit_log(days)
    entries = _filter_memory_audit_entries(all_entries, op=op, category=category)
    period = _memory_review_period(days)
    if not entries:
        filters = []
        if op:
            filters.append(f"op={op}")
        if category:
            filters.append(f"category={category}")
        filter_text = f" with filters ({', '.join(filters)})" if filters else ""
        return f"📋 Memory Review: {period.lower()}{filter_text} no records exist."

    # Grouping by operation
    adds        = [e for e in entries if e.get("op") == "add"]
    overwrites  = [e for e in entries if e.get("op") == "overwrite"]
    alongside   = [e for e in entries if e.get("op") == "add_alongside"]
    skip_dup    = [e for e in entries if e.get("op") == "skip_duplicate"]
    skip_old    = [e for e in entries if e.get("op") == "skip_keep_old"]
    reflections = [e for e in entries if e.get("op") in ("reflection_applied", "reflection_saved")]

    filters = []
    if op:
        filters.append(f"op={op}")
    if category:
        filters.append(f"category={category}")
    filter_text = f" ({', '.join(filters)})" if filters else ""
    has_filter = _memory_review_has_filter(op=op, category=category)
    lines = [f"📋 *Memory Review — {period}{filter_text}: {len(entries)} memory actions*\n"]

    _append_memory_review_section(
        lines,
        title="✅ *Learned / kept new*",
        entries=adds,
        has_filter=has_filter,
        limit=5,
        formatter=lambda e: f"  [{e.get('ts','')}] [{e.get('category','?')}] {e.get('fact','')[:80]}",
    )
    _append_memory_review_section(
        lines,
        title="♻️ *Corrected old memories*",
        entries=overwrites,
        has_filter=has_filter,
        limit=5,
        formatter=lambda e: f"  [{e.get('ts','')}] {e.get('fact','')[:60]} ← {e.get('old','')[:40]} ({e.get('reason','')})",
    )
    _append_memory_review_section(
        lines,
        title="🧩 *Kept similar memories separate*",
        entries=alongside,
        has_filter=has_filter,
        limit=5,
        formatter=lambda e: f"  [{e.get('ts','')}] {e.get('fact','')[:70]} (dist={e.get('distance','?')}, overlap={e.get('overlap','?')})",
    )
    _append_memory_review_section(
        lines,
        title="🔁 *Ignored as duplicates*",
        entries=skip_dup,
        has_filter=has_filter,
        limit=3,
        formatter=lambda e: f"  [{e.get('ts','')}] {e.get('fact','')[:70]} (dist={e.get('distance','?')})",
    )
    _append_memory_review_section(
        lines,
        title="🔒 *Kept older/richer memory*",
        entries=skip_old,
        has_filter=has_filter,
        limit=3,
        formatter=lambda e: f"  [{e.get('ts','')}] {e.get('fact','')[:70]}",
    )
    _append_memory_review_section(
        lines,
        title="🧠 *Lessons / reflections*",
        entries=reflections,
        has_filter=has_filter,
        limit=5,
        formatter=lambda e: f"  [{e.get('ts','')}] {'✓ applied' if e.get('op') == 'reflection_applied' else 'saved'}: {(e.get('lesson') or e.get('observation') or '')[:80]}",
    )

    return "\n".join(lines)


all_tools = [
    search_memory, save_to_memory, delete_from_memory, retrieve_photo, update_pending_linkedin_post, process_and_clear_linkedin_post,
    set_local_reminder, manage_list,
    google_calendar_tool, google_tasks_tool, drive_manager,
    read_local_file, write_code, run_code, write_custom_tool,
    mail_manager, github_manager, control_vacuum, control_spotify, recipe_expert, log_meal, search_recipe_library, get_saved_recipe, mark_recipe_favorite, search_flights, search_google_places,
    create_file_tool, get_current_location,
    get_news, get_weather_forecast, search_supermarket_prices, relay_local_payload,
    search_goldmall_offers, execute_local_pipeline, archive_file, get_navigation_info, generate_image_tool, post_to_linkedin, learn_routine, edit_routine, delete_routine, get_routines, search_routines, control_routine_notifications, control_routine_schedule, control_routine_condition, control_routine_cooldown, control_pending_followup, browse_url,
    duckduckgo_search, run_terminal_command, get_fit_summary, save_goal_tool, update_goal_status_tool, update_goal_progress_tool, update_goal_milestones_tool, tool_stats, system_doctor, memory_review,
    repo_mapper,
    scan_receipt,
    text_stats,
    register_tool,
    research_last30days, morning_briefing,
    hn_briefing,
    # Project tools
    grant_project_access, list_project_files, read_project_file,
    edit_project_file, write_project_file, grep_project_files,
    list_recent_files,
    # File generator
    generate_excel, generate_word_doc, generate_pdf, generate_csv,
    list_agent_skills, read_agent_skill, run_officecli, manage_context_flag,
    get_world_time,
]
